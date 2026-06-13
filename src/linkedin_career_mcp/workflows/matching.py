from __future__ import annotations

import argparse
import asyncio
import json
import math
import re
import sys
from collections import Counter
from collections.abc import Callable, Mapping
from dataclasses import dataclass, field
from datetime import UTC, date, datetime
from html import escape
from pathlib import Path
from typing import Any, Literal

from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.worksheet import Worksheet
from pydantic import BaseModel, ConfigDict, Field
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer

from linkedin_career_mcp.api_client import ApiLlmClient
from linkedin_career_mcp.ats import AtsProxyScore, calculate_ats_proxy_score
from linkedin_career_mcp.config import Settings, load_settings
from linkedin_career_mcp.errors import LinkedInCareerMcpError, WorkflowError
from linkedin_career_mcp.models import DatePosted, JobDetails, JobPosting, JobSearchQuery
from linkedin_career_mcp.ollama import OllamaClient
from linkedin_career_mcp.providers import LinkedInPublicJobsProvider
from linkedin_career_mcp.query_optimizer import (
    QueryOutcome,
    ScoredQuery,
    StoredQueryOutcome,
    historical_query_candidates,
    load_query_outcomes,
    rank_search_queries,
    record_query_outcome,
)
from linkedin_career_mcp.services import JobSearchService
from linkedin_career_mcp.webapp import (
    DEFAULT_DATABASE as APPLICATION_DATABASE,
)
from linkedin_career_mcp.webapp import (
    ApplicationJobRecord,
    connect_database,
    fetch_application_job_ids,
    fetch_application_job_records,
    fetch_existing_cover_letter_job_ids,
    fetch_existing_resume_job_ids,
    import_output_artifacts,
    upsert_application_artifact,
)

DEFAULT_PROFILE_DIR = Path("profile")
DEFAULT_BLACKLIST_PATH = Path(".blacklist")
DEFAULT_OUTPUT_DIR = Path("output")
DEFAULT_SOURCE_RESUME = "MP-RESUME-AGENTIC.pdf"
DEFAULT_CURRENT_JOB_DESCRIPTION = "Senior_Platform_Software_Engineer(IC3).pdf"
TRACKING_WORKBOOK = Path("tracking/read_applications/linkedin_applications.xlsx")
SUPPORTED_TEXT_SUFFIXES = {".csv", ".json", ".md", ".rst", ".text", ".txt"}
SUPPORTED_PROFILE_SUFFIXES = SUPPORTED_TEXT_SUFFIXES | {".docx", ".pdf"}
ArtifactMode = Literal["all", "resumes-only", "cover-letters-only"]
ProgressCallback = Callable[[str], None]
DISALLOWED_EXPERIENCE_LEVELS = {"internship", "entry_level"}
SEARCH_EXPERIENCE_LEVELS = ("associate", "mid_senior", "director", "executive")
NO_PUBLIC_JOB_DESCRIPTION = "No public job description was available."
RESUME_HEADER_NAME = "Max Perkhounkov"
LINKEDIN_PROFILE_LABEL = "linkedin.com/in/maxim-perkhounkov"
LINKEDIN_PROFILE_URL = "https://www.linkedin.com/in/maxim-perkhounkov/"
LINKEDIN_PROFILE_MARKDOWN = f"[{LINKEDIN_PROFILE_LABEL}]({LINKEDIN_PROFILE_URL})"
RESUME_HEADER_CONTACT = (
    "Iowa City, IA | 641-781-0477 | mperkhounkov1@gmail.com | "
    f"{LINKEDIN_PROFILE_MARKDOWN}"
)
EMERALD_ACCENT = colors.HexColor("#57BA86")
EMERALD_DARK = colors.HexColor("#047857")
RESUME_BODY_COLOR = colors.HexColor("#111827")
RESUME_MUTED_COLOR = colors.HexColor("#475569")
STATIC_PROFESSIONAL_SUMMARY = (
    "Analytical and metrics-driven Senior Platform Software Engineer with over 10 years of "
    "multi-disciplinary experience architecting scalable distributed systems, developer tooling, "
    "and cloud automation frameworks. Proven track record leading enterprise-level integrations, "
    "optimizing platform resilience, and implementing secure API and observability pipelines "
    "across 40,000+ devices. Combines a strong background in advanced mathematics and "
    "algorithmic problem-solving with hands-on expertise in CI/CD, Infrastructure as Code (IaC), "
    "and modern multi-tenant cloud architectures."
)
AI_GENERATION_NOTE = (
    "Note: This resume is custom tailored for every job position using my automated agentic "
    "workflow found at: "
    "[mperkhou/linkedin-career-mcp](https://github.com/mperkhou/linkedin-career-mcp)"
)
COVER_LETTER_ORACLE_OPENER = (
    "In my current role at Oracle, I work as a Senior Technical Lead and Cloud Automation "
    "Engineer building and maintaining large-scale platform automation systems across "
    "international datacenters."
)
COVER_LETTER_PRIOR_EXPERIENCE_OPENER = (
    "My earlier experience also strengthens my fit for this position."
)
COVER_LETTER_PROJECT_PARAGRAPH = (
    "I also want to highlight the automation project "
    "([mperkhou/linkedin-career-mcp](https://github.com/mperkhou/linkedin-career-mcp)) used "
    "to generate the resume and cover letter submitted with this application. I built a custom "
    "agentic workflow that searches public LinkedIn job postings, compares each Job Opening "
    "Description against my resume and current role context, and generates tailored resume "
    "artifacts through cost-conscious API calls to OpenRouter and DeepSeek. I developed the "
    "project using multiple AI-assisted engineering tools, including Codex, Cline with DeepSeek, "
    "and GitHub Copilot, while actively managing prompt structure, context windows, token usage, "
    "model selection, and output validation. I built this in my own time because I genuinely "
    "enjoy automation, AI tooling, and turning repetitive workflows into reliable systems."
)
COVER_LETTER_DEFAULT_OPENING_ALIGNMENT = (
    "the platform engineering, automation, distributed systems, and AI-focused capabilities "
    "you are looking for"
)
COVER_LETTER_DEFAULT_ORACLE_ALIGNMENT = (
    "I have owned multi-tenant infrastructure components serving 40,000+ managed endpoints, "
    "built distributed observability pipelines using Filebeat, Logstash, and OpenSearch, and "
    "developed Python automation frameworks to replace legacy third-party tooling. I also "
    "standardize platform services with Terraform, Chef, Ansible, and CI/CD pipelines so "
    "operational work stays repeatable across large enterprise environments. That experience "
    "maps directly to roles requiring resilient platform services, secure API integrations, "
    "debugging across distributed systems, developer tooling, CI/CD, and practical automation "
    "at enterprise scale."
)
COVER_LETTER_DEFAULT_PRIOR_EXPERIENCE_ALIGNMENT = (
    "At the University of Iowa Hospitals and Clinics, I built Python and AutoIT automation "
    "scripts, supported regulated medical platforms, and contributed to a React-based DICOM "
    "anonymization server. At Steindler Orthopedic Clinic, Stamats Communications, and VIDA "
    "Diagnostics, I led infrastructure modernization, Azure migrations, monitoring and reporting "
    "workflows, Python data-transfer tooling, and a Django platform rebuild. Across those roles, "
    "I developed the mix of software engineering, systems thinking, data movement, cloud "
    "infrastructure, and cross-functional troubleshooting needed to contribute effectively in "
    "complex technical environments."
)
RESUME_SECTION_HEADINGS = {
    "Professional Summary",
    "Core Technical Skills",
    "Professional Experience",
    "Education & Certifications",
}
JOB_DESCRIPTION_PROMPT_MAX_CHARS = 12_000
ATS_REPAIR_TARGET_SCORE = 90
ATS_REPAIR_MAX_ATTEMPTS = 2
ATS_REPAIR_MAX_MISSING_TERMS = 8
ATS_REPAIR_MAX_SOURCE_SNIPPETS = 8
ATS_REPAIR_SOURCE_SNIPPET_RADIUS = 1
SOURCE_EVIDENCE_TERM_ALIASES = {
    "artificialintelligence": ("ai", "agenticai", "aiassisted", "aitooling", "appliedai"),
    "developerproductivity": (
        "developertooling",
        "developerworkflow",
        "developerworkflows",
        "deliveryvelocity",
    ),
    "developerproductivitycicd": (
        "developertooling",
        "developerworkflow",
        "developerworkflows",
        "deliveryvelocity",
        "cicd",
    ),
    "productivitycicd": (
        "developertooling",
        "developerworkflow",
        "developerworkflows",
        "deliveryvelocity",
        "cicd",
    ),
}
DISALLOWED_CORE_SKILLS = {
    "error budget",
    "error budgets",
    "error budget analysis",
}
LEADING_BULLET_RE = re.compile(
    r"^(?:[-*]|\u2022|\u2023|\u25e6|\u2043|\u2219|\u25cf|\u25cb)\s+"
)
ROLE_RELEVANT_START_HEADINGS = (
    "Job Summary",
    "Position Summary",
    "Role Summary",
    "The Role",
    "About the Role",
    "About this Role",
    "About the Job",
    "The Opportunity",
    "What You’ll Do",
    "What You'll Do",
    "What You Will Do",
    "What You’ll Be Doing",
    "What You'll Be Doing",
    "What You Will Be Doing",
    "What We Need To See",
    "Key Responsibilities",
    "Responsibilities",
    "What You’ll Bring",
    "What You'll Bring",
    "What You Bring",
    "What We’re Looking For",
    "What We're Looking For",
    "Required Qualifications",
    "Minimum Qualifications",
    "Basic Qualifications",
    "Qualifications",
    "Requirements",
    "Skills and Experience",
    "Who You Are",
    "You might thrive",
    "You could be",
)
LOW_SIGNAL_PREAMBLE_HEADINGS = (
    "Our Mission",
    "Our Mission & Values",
    "Mission & Values",
    "Our Values",
    "Our Culture",
    "Our Culture & Work Style",
    "Culture & Work Style",
    "Life at",
    "Why Join",
    "Why Join Us",
    "Why Join The",
)
TRAILING_BOILERPLATE_HEADINGS = (
    "How We Support You",
    "Why Join",
    "Why Join Us",
    "Why You’ll Love Working Here",
    "Why You'll Love Working Here",
    "Benefits & Perks",
    "Pay & Benefits",
    "Our Benefits",
    "Perks & Benefits",
    "Health & Wellness",
    "Financial Well-being",
    "Family Support",
    "Growth & Development",
    "Time Off & Flexibility",
    "What We Offer",
    "Compensation",
    "Compensation Range",
    "Equal Opportunity",
    "Equal opportunity employer",
    "Diversity, Equity",
    "How we feel about Diversity",
    "Accommodations",
    "For US Applicants",
    "Benefits Offering",
    "Privacy Statement",
    "Privacy Notice",
    "Applicant Privacy Notice",
    "Applicant Notice",
    "By providing your information",
    "Your base salary",
    "US base salary range",
    "US Salary Range",
    "The base salary range",
    "Base salary range",
    "Salary range",
    "Applications for this job",
    "This posting is for",
    "NVIDIA uses AI tools",
    "About The Team",
)
ROLE_RELEVANT_PREFIX_MARKERS = (
    "you will",
    "you'll",
    "responsibil",
    "build",
    "design",
    "develop",
    "automate",
    "platform",
    "infrastructure",
    "cloud",
    "api",
    "monitor",
    "observability",
    "terraform",
    "kubernetes",
    "aws",
    "python",
    "ci/cd",
    "sagemaker",
    "snowflake",
    "splunk",
    "software engineering expectations",
    "developer productivity",
)
JOD_HARD_DROP_MARKERS = (
    "base salary",
    "salary range",
    "compensation",
    "competitive compensation",
    "target incentive compensation",
    "on-target-earnings",
    "internal pay equity",
    "benefits",
    "wellness benefits",
    "pay transparency",
    "medical, dental",
    "401(k)",
    "parental leave",
    "equal opportunity",
    "equal employment opportunity",
    "affirmative action",
    "reasonable accommodation",
    "accommodation requests",
    "privacy policy",
    "applicant privacy",
    "personal information",
    "e-verify",
    "recruiter will share",
    "hiring process",
    "interview process",
    "application process",
    "starting pay",
    "we may use artificial intelligence",
    "we may use ai tools",
    "401k",
    "paid company holidays",
)
JOD_HARD_KEEP_MARKERS = (
    "responsibil",
    "what you'll do",
    "what you will do",
    "key responsibilities",
    "job summary",
    "role summary",
    "required qualifications",
    "minimum qualifications",
    "basic qualifications",
    "skills and experience",
)
JOD_CHUNK_KEEP_THRESHOLD = -0.75
JOD_CHUNK_MIN_MEANINGFUL_LENGTH = 24
INLINE_TRAILING_BOILERPLATE_BOUNDARY_RE = re.compile(
    r"\s+(?=(?:"
    r"Benefits\s+Compensation|"
    r"Benefits\b|"
    r"Salary\s+Ranges?|"
    r"Salary\s+&\s+Benefits|"
    r"Pay\s+Disclai\\s*mer|"
    r"Pay\s+Range|"
    r"In\s+the\s+spirit\s+of\s+pay\s+transparency|"
    r"Actual\s+placement\s+in\s+range|"
    r"What\s+[A-Z][A-Za-z0-9]+\s+Offers\s+You|"
    r"Competitive\s+compensation|"
    r"To\s+determine\s+a\s+successful\s+candidate|"
    r"It\s+is\s+the\s+policy\s+of|"
    r"Accommodation\s+requests|"
    r"Disclai\s*me\s*r|"
    r"EE\s*O\b|"
    r"We\s+offer\s+a\s+401k|"
    r"Additional\s+factors\s+considered|"
    r"Actual\s+compensation|"
    r"Individual\s+total\s+compensation|"
    r"The\s+base\s+compensation\s+range|"
    r"Compensation\s+(?:The\s+salary\s+range|for|of|&\s+Benefits|\$)|"
    r"Compensation\b"
    r")\b)",
)
JOD_BOILERPLATE_START_RE = re.compile(
    r"^(?:"
    r"actual compensation|"
    r"actual placement in range|"
    r"benefits\b|"
    r"competitive compensation|"
    r"compensation\b|"
    r"in the spirit of pay transparency|"
    r"individual total compensation|"
    r"it is the policy of .{0,80}?equal employment opportunity|"
    r"to determine a successful candidate|"
    r"we offer a 401k|"
    r"the base compensation range|"
    r"the opportunity to work alongside .{0,160}?compensation|"
    r"what [a-z0-9]+ offers you|"
    r"[a-z0-9 .,'&-]{0,80}? is an equal opportunity employer"
    r")",
    re.IGNORECASE,
)
JOD_CHUNK_TRAINING_EXAMPLES = (
    (
        "keep",
        "Job Summary Build scalable platform services, APIs, automation, observability, "
        "and infrastructure used by engineering teams.",
    ),
    (
        "keep",
        "Responsibilities Design and develop cloud automation, CI/CD pipelines, monitoring, "
        "and distributed systems for production workloads.",
    ),
    (
        "keep",
        "Required Qualifications Experience with Python, Terraform, AWS, Kubernetes, Linux, "
        "security, APIs, and operational troubleshooting.",
    ),
    (
        "keep",
        "What You'll Do Own developer tooling, reliability, data pipelines, integrations, "
        "and automation frameworks across teams.",
    ),
    (
        "keep",
        "Preferred Qualifications Experience with LLM workflows, AI tooling, platform "
        "engineering, observability, and infrastructure as code.",
    ),
    (
        "keep",
        "The elevator pitch is owning automation that makes manual processes disappear "
        "through Python tooling, endpoint management, monitoring, and permanent systems fixes.",
    ),
    (
        "keep",
        "Hammerspace delivers a global data environment spanning data centers, Linux systems, "
        "AWS, Azure, Google cloud infrastructure, networking, storage, and distributed systems.",
    ),
    (
        "keep",
        "Modern AI workloads place different demands on infrastructure, including LLM inference, "
        "agent runtimes, vector stores, GPU economics, and real-time computer vision pipelines.",
    ),
    (
        "drop",
        "Benefits include comprehensive medical dental and vision coverage, wellness "
        "stipends, parental leave, paid time off, and retirement savings.",
    ),
    (
        "drop",
        "The base salary range for this role depends on location, market data, equity, "
        "bonus eligibility, and other compensation factors.",
    ),
    (
        "drop",
        "We are an equal opportunity employer and provide reasonable accommodation during "
        "the application and interview process.",
    ),
    (
        "drop",
        "Applicant privacy notice explains how personal information is processed, retained, "
        "and shared under our privacy policy.",
    ),
    (
        "drop",
        "A recruiter will share more details about hiring process logistics, interview "
        "steps, application review, and employment eligibility.",
    ),
    (
        "drop",
        "Actual compensation awarded to successful candidates is based on market, location, "
        "scope, individual qualifications, and interview process assessment.",
    ),
    (
        "drop",
        "Individual total compensation will vary based on qualifications, skill level, "
        "competencies, role location, salary range, and work location.",
    ),
    (
        "drop",
        "The opportunity includes a competitive compensation package, 401(k), health, dental, "
        "vision coverage, paid vacation days, office perks, and social events.",
    ),
)
DEFAULT_CORE_TECHNICAL_SKILLS: tuple[tuple[str, tuple[str, ...]], ...] = (
    (
        "Languages & Frameworks",
        (
            "Python",
            "Ruby",
            "JavaScript",
            "Node.js",
            "React.js",
            "Go",
            "Bash",
            "PowerShell",
            "Ruby on Rails",
            "Django",
        ),
    ),
    (
        "Distributed Systems & Cloud",
        (
            "AWS",
            "Azure",
            "Oracle Cloud Infrastructure (OCI)",
            "ElasticSearch",
            "OpenSearch",
        ),
    ),
    (
        "Platform & API Engineering",
        (
            "RESTful APIs",
            "Systems Architecture",
            "Microservices",
            "JSON/XML",
            "API Integration",
        ),
    ),
    (
        "Automation & IaC",
        (
            "Chef (Cookbooks/Policies)",
            "Ansible",
            "Terraform",
            "Jenkins",
            "CloudLab CI/CD Pipelines",
        ),
    ),
    (
        "Data & Observability",
        (
            "Filebeat",
            "Logstash",
            "PostgreSQL",
            "MongoDB",
            "SQL",
            "Data Pipelines",
            "Observability Dashboards",
        ),
    ),
    (
        "Security & Compliance",
        (
            "Secure Coding Practices",
            "Vulnerability Mitigation",
            "Role-Based Access Control (RBAC)",
        ),
    ),
    (
        "AI Tools",
        (
            "Codex",
            "Oracle Code Assist (OCA)",
            "Cline",
            "OpenRouter",
            "ChatGPT",
            "LLM Prompting",
        ),
    ),
)
DEFAULT_SUPPLEMENTAL_SEARCH_KEYWORDS = (
    "Senior Platform Engineer AI",
    "Cloud Automation Engineer LLMs",
    "Infrastructure Software Engineer agentic AI",
    "DevOps Engineer distributed systems",
)
LEGACY_TRACKING_HEADERS = [
    "job_id",
    "company",
    "job_title",
    "linkedin_url",
    "customized_resume",
    "applied_to",
    "date_applied",
]
TRACKING_HEADERS = [
    "job_id",
    "company",
    "job_title",
    "linkedin_url",
    "customized_resume",
    "cover_letter",
    "applied_to",
    "date_applied",
]
DEFAULT_SCJDIR = """
Oracle | Remote / International Datacenters
Senior Technical Lead - Cloud Automation Engineer | Feb 2022 - Present
Platform Component Ownership: Architected and managed the end-to-end multi-tenant architecture
of a global Chef infrastructure orchestrating contracts for 40,000+ managed endpoints across
international datacenters.
Distributed Observability: Built and scaled an enterprise data pipeline utilizing Filebeat agents
on 40,000+ devices routing via Logstash to centralized OpenSearch clusters; analyzed usage, logs,
and operational trends to optimize reliability.
Integration & API Frameworks: Developed a custom Python package and playbooks framework within
Oracle Linux Automation Manager (OLAM) to replace a legacy third-party platform; unified and
automated cross-vendor API integrations for 12,000+ network surfaces.
Infrastructure as Code (IaC): Engineered reusable, highly available Terraform plans, Chef
Cookbooks, and Ansible roles to standardize platform SDKs and services across all enterprise
domains.
CI/CD & Resilience: Eliminated production configuration drift and boosted delivery velocity by
replacing manual workflows with automated Jenkins and CloudLab CI/CD release pipelines.
Developer Tooling Innovation: Spearheaded team-level adoption of AI-assisted engineering tools
(Cline, Codex, Oracle Code Assist), developing reliable internal workflows that reduced
test-driven development (TDD) busywork by 80%.
""".strip()
DEFAULT_PRIOR_EXPERIENCE_ENTRIES: tuple[dict[str, object], ...] = (
    {
        "organization": "University of Iowa Hospitals and Clinics",
        "location": "Iowa City, IA",
        "title": "Engineering Support Specialist",
        "dates": "Jan 2020 - May 2021",
        "bullets": (
            "Full-Cycle Software Engineering: Adhered to strict software development lifecycles "
            "to build custom Python and AutoIT automation scripts, streamlining system upgrades "
            "across hundreds of mission-critical platform nodes.",
            "Web Application Development: Collaborated on the structural design and "
            "implementation of a DICOM anonymization server utilizing a modern React.js frontend "
            "interface.",
            "Deep Debugging & Patching: Conducted performance troubleshooting, defect handling, "
            "and remote patch deployments on highly regulated medical platform surfaces.",
        ),
    },
    {
        "organization": "Steindler Orthopedic Clinic",
        "location": "Iowa City, IA",
        "title": "IT Administrator / Systems Engineer",
        "dates": "Mar 2019 - Nov 2019",
        "bullets": (
            "Cloud Migrations & Architecture: Led the engineering lifecycle to modernize 7+ year "
            "old core infrastructure systems, executing legacy virtualization overhauls via "
            "ESX/VMware and migrating services to Azure Cloud.",
            "API & Workflow Integration: Built custom PHP plugins and integrated secure Azure "
            "SharePoint document-control workflows to boost internal process cross-functional "
            "alignment.",
            "Observability Dashboards: Launched an enterprise ticketing and incident-tracking "
            "system, designing centralized health monitoring dashboards and automated analytical "
            "reporting tools.",
        ),
    },
    {
        "organization": "Stamats Communications",
        "location": "Cedar Rapids, IA",
        "title": "Systems Administrator (Contract)",
        "dates": "Apr 2018 - Oct 2018",
        "bullets": (
            "Infrastructure Optimization: Partnered with cross-functional leadership to architect "
            "and execute a multi-million dollar infrastructure upgrade, maximizing capacity and "
            "network availability.",
            "Cloud Ecosystem Deployment: Spearheaded on-premise Exchange migrations to Azure "
            "cloud environments while ensuring strict alignment with enterprise information "
            "security standards.",
        ),
    },
    {
        "organization": "VIDA Diagnostics",
        "location": "Coralville, IA",
        "title": "Systems Engineer",
        "dates": "Mar 2014 - Mar 2018",
        "bullets": (
            "Algorithmic Data Engineering: Wrote highly scalable Python data-transfer scripts "
            "optimizing the ingestion and transit of massive CT imagery cache structures between "
            "Linux environments and SAN/NAS storage arrays.",
            "Framework Refactoring: Rebuilt corporate web platforms entirely from WordPress to a "
            "robust, secure Django framework to improve backend structural integrity.",
        ),
    },
)
DEFAULT_EDUCATION_CERTIFICATIONS = (
    "- Bachelor of Science in Physics & Mathematics | University of Iowa, IA",
    "  - Focus: Graduate-level mathematics, applied statistics, and computer science principles.",
    (
        "  - Leadership: Teaching Assistant (Physics Department), President of the University "
        "Chess Club."
    ),
    "- Oracle Cloud Infrastructure (OCI) Engineer | Certification (August, 2024)",
    "- Oracle Cloud Infrastructure AI Foundations Associate | Certification (May, 2026)",
    "- AlienVault Certified Security Engineer (AVCSE) | Certification",
    "- Advanced Continuing Education (Udemy): Docker & Kubernetes Ecosystems, Microservices "
    "Engineering (Node.js & React), PostgreSQL Database Bootcamp, Object-Oriented Programming "
    "(OOP) & Agile Methodologies.",
)


@dataclass(frozen=True)
class ProfileDocument:
    path: Path
    text: str


class TailoredResumeArtifact(BaseModel):
    model_config = ConfigDict(extra="forbid")

    job_id: str
    company: str | None
    title: str
    linkedin_url: str | None
    resume_path: str = ""
    cover_letter_path: str | None = None
    artifact_kind: Literal["resume", "recommendations", "cover_letter"] = "resume"
    recommendations_path: str | None = None


class MissingApplicationArtifact(BaseModel):
    model_config = ConfigDict(extra="forbid")

    job_id: str
    company: str
    title: str
    missing: list[Literal["resume", "cover_letter"]]


class ArtifactAudit(BaseModel):
    model_config = ConfigDict(extra="forbid")

    total_jobs: int = 0
    with_resumes: int = 0
    with_cover_letters: int = 0
    missing_resumes: int = 0
    missing_cover_letters: int = 0
    missing_artifacts: list[MissingApplicationArtifact] = Field(default_factory=list)


class CoverLetterRetryRecord(BaseModel):
    model_config = ConfigDict(extra="forbid")

    attempt: int
    job_id: str
    company: str | None
    title: str
    status: Literal["created", "failed"]
    cover_letter_path: str | None = None
    error: str | None = None


class MatchingJobsWorkflowResult(BaseModel):
    model_config = ConfigDict(extra="forbid")

    profile_files: list[str]
    search_queries: list[JobSearchQuery]
    jobs_found: int
    resumes_created: int
    cover_letters_created: int = 0
    recommendations_created: int = 0
    tracking_spreadsheet: str
    skipped_blacklisted: list[str] = Field(default_factory=list)
    skipped_existing: list[str] = Field(default_factory=list)
    skipped_workplace_type: list[str] = Field(default_factory=list)
    skipped_experience_level: list[str] = Field(default_factory=list)
    errors: list[str] = Field(default_factory=list)
    artifacts: list[TailoredResumeArtifact] = Field(default_factory=list)
    artifact_audit: ArtifactAudit = Field(default_factory=ArtifactAudit)
    cover_letter_retries: list[CoverLetterRetryRecord] = Field(default_factory=list)


class CompanyBlacklist:
    def __init__(self, patterns: list[str]) -> None:
        self._patterns = [pattern.strip() for pattern in patterns if pattern.strip()]

    @classmethod
    def from_file(cls, path: Path) -> CompanyBlacklist:
        if not path.exists():
            return cls([])
        patterns: list[str] = []
        for line in path.read_text(encoding="utf-8").splitlines():
            value = line.strip()
            if not value or value.startswith("#"):
                continue
            patterns.append(value)
        return cls(patterns)

    def matches(self, company: str | None) -> bool:
        if not company:
            return False
        company_value = company.casefold()
        return any(_glob_matches(company_value, pattern.casefold()) for pattern in self._patterns)


@dataclass
class _SearchMemory:
    """Tracks which keyword groups returned jobs and which did not, so the LLM can
    refine successive iterations."""

    rewarded_keywords: list[str] = field(default_factory=list)
    penalized_keywords: list[str] = field(default_factory=list)
    attempted_query_keys: set[str] = field(default_factory=set)
    total_searches: int = 0

    def query_key(self, query: JobSearchQuery) -> str:
        wt = query.workplace_type or "any"
        return f"{query.keywords.casefold()}::{query.location.casefold()}::{wt}"

    def register_result(self, query: JobSearchQuery, count: int) -> None:
        self.total_searches += 1
        self.attempted_query_keys.add(self.query_key(query))
        if count > 0:
            self.rewarded_keywords.append(query.keywords)
        else:
            self.penalized_keywords.append(query.keywords)

    @property
    def reward_sample(self) -> str:
        if not self.rewarded_keywords:
            return "No keywords have produced results yet."
        return "\n".join(
            f"- {kw}" for kw in self.rewarded_keywords[-10:]
        )

    @property
    def penalty_sample(self) -> str:
        if not self.penalized_keywords:
            return "None so far."
        return "\n".join(
            f"- {kw}" for kw in self.penalized_keywords[-10:]
        )

    def has_query(self, query: JobSearchQuery) -> bool:
        return self.query_key(query) in self.attempted_query_keys


@dataclass(frozen=True)
class _RegenerationCandidate:
    job: JobDetails
    stored_job_description: str | None


@dataclass(frozen=True)
class _ResumeDraft:
    tailored_scjdir: str
    sections_plan: Mapping[str, Any]
    text: str


@dataclass
class _QueryRunOutcome:
    scored_query: ScoredQuery
    results_returned: int = 0
    fresh_jobs_accepted: int = 0
    skipped_existing: int = 0
    skipped_blacklisted: int = 0
    skipped_workplace_type: int = 0
    skipped_experience_level: int = 0
    resumes_generated: int = 0
    ats_scores: list[int] = field(default_factory=list)

    @property
    def query(self) -> JobSearchQuery:
        return self.scored_query.query

    def to_query_outcome(self, *, artifact_mode: ArtifactMode) -> QueryOutcome:
        average_ats_score = (
            sum(self.ats_scores) / len(self.ats_scores)
            if self.ats_scores
            else None
        )
        return QueryOutcome(
            query=self.query,
            profile_match=self.scored_query.profile_match,
            query_score=self.scored_query.score,
            results_returned=self.results_returned,
            fresh_jobs_accepted=self.fresh_jobs_accepted,
            skipped_existing=self.skipped_existing,
            skipped_blacklisted=self.skipped_blacklisted,
            skipped_workplace_type=self.skipped_workplace_type,
            skipped_experience_level=self.skipped_experience_level,
            resumes_generated=self.resumes_generated,
            average_ats_score=average_ats_score,
            artifact_mode=artifact_mode,
        )


MAX_ITERATIVE_SEARCHES = 1000
MIN_SEARCHES_BEFORE_STOP = 4


class MatchingJobsWorkflow:
    def __init__(
        self,
        *,
        service: JobSearchService,
        ollama: Any,
        planner_llm: Any | None = None,
    ) -> None:
        self._service = service
        self._ollama = ollama
        self._planner_llm = planner_llm or ollama

    async def run(
        self,
        *,
        profile_dir: Path = DEFAULT_PROFILE_DIR,
        blacklist_path: Path = DEFAULT_BLACKLIST_PATH,
        output_dir: Path = DEFAULT_OUTPUT_DIR,
        source_resume_name: str = DEFAULT_SOURCE_RESUME,
        current_job_description_name: str = DEFAULT_CURRENT_JOB_DESCRIPTION,
        location: str = "United States",
        date_posted: DatePosted = "past_week",
        limit_per_query: int = 10,
        max_queries: int = 6,
        max_jobs: int = 10,
        artifact_mode: ArtifactMode = "all",
        cover_letter_retry_attempts: int = 0,
        progress_callback: ProgressCallback | None = None,
    ) -> MatchingJobsWorkflowResult:
        profile_documents = load_profile_documents(profile_dir)
        profile_context = format_profile_context(profile_documents)
        source_resume = _find_source_resume(profile_documents, source_resume_name)
        current_job_description = _find_source_resume(
            profile_documents,
            current_job_description_name,
        )
        blacklist = CompanyBlacklist.from_file(blacklist_path)

        candidates: list[JobDetails] = []
        skipped_blacklisted: list[str] = []
        skipped_existing: list[str] = []
        skipped_workplace_type: list[str] = []
        skipped_experience_level: list[str] = []
        errors: list[str] = []
        seen_job_ids: set[str] = set()
        tracking_path = output_dir / TRACKING_WORKBOOK
        application_database_path = output_dir / APPLICATION_DATABASE
        if tracking_path.exists():
            import_output_artifacts(output_dir=output_dir, database_path=application_database_path)
        existing_resume_job_ids = fetch_existing_resume_job_ids(application_database_path)
        existing_cover_letter_job_ids = fetch_existing_cover_letter_job_ids(
            application_database_path
        )
        existing_application_job_ids = fetch_application_job_ids(application_database_path)
        existing_artifact_job_ids = (
            existing_cover_letter_job_ids
            if artifact_mode == "cover-letters-only"
            else existing_resume_job_ids
        )
        query_history = load_query_outcomes(application_database_path)
        search_memory = _SearchMemory()
        all_search_queries: list[JobSearchQuery] = []
        query_run_outcomes: list[_QueryRunOutcome] = []
        accepted_job_outcomes: dict[str, _QueryRunOutcome] = {}
        min_searches_before_stop = min(max(max_queries, 1), MIN_SEARCHES_BEFORE_STOP)

        while (
            len(candidates) < max_jobs
            or search_memory.total_searches < min_searches_before_stop
        ) and search_memory.total_searches < MAX_ITERATIVE_SEARCHES:
            scored_search_queries = await self._generate_search_queries(
                profile_context=profile_context,
                location=location,
                date_posted=date_posted,
                limit_per_query=limit_per_query,
                max_queries=max_queries,
                query_history=query_history,
                search_memory=search_memory,
            )

            new_query_found = False
            for scored_query in scored_search_queries:
                query = scored_query.query
                if search_memory.total_searches >= MAX_ITERATIVE_SEARCHES:
                    break
                if search_memory.has_query(query):
                    continue
                new_query_found = True
                all_search_queries.append(query)
                query_outcome = _QueryRunOutcome(scored_query=scored_query)
                query_run_outcomes.append(query_outcome)
                excluded_job_ids = existing_artifact_job_ids | seen_job_ids
                search_query = query.model_copy(update={"exclude_job_ids": excluded_job_ids})
                try:
                    result = await self._service.search(search_query)
                except LinkedInCareerMcpError as exc:
                    errors.append(f"{query.keywords}: {exc}")
                    search_memory.register_result(query, 0)
                    continue
                search_memory.register_result(query, len(result.jobs))
                query_outcome.results_returned = len(result.jobs)
                for posting in result.jobs:
                    if len(candidates) >= max_jobs:
                        break
                    if posting.job_id in seen_job_ids:
                        continue
                    seen_job_ids.add(posting.job_id)
                    if posting.job_id in existing_artifact_job_ids:
                        skipped_existing.append(_job_label(posting.company, posting.title))
                        query_outcome.skipped_existing += 1
                        continue
                    if blacklist.matches(posting.company):
                        skipped_blacklisted.append(_job_label(posting.company, posting.title))
                        query_outcome.skipped_blacklisted += 1
                        continue
                    try:
                        details = await self._service.get_details(
                            str(posting.job_url or posting.job_id),
                        )
                    except LinkedInCareerMcpError as exc:
                        errors.append(f"{posting.job_id}: {exc}")
                        details = JobDetails(**posting.model_dump())
                    details = _merge_posting_with_fetched_details(posting, details)
                    if blacklist.matches(details.company):
                        skipped_blacklisted.append(_job_label(details.company, details.title))
                        query_outcome.skipped_blacklisted += 1
                        continue
                    if not _is_remote_or_hybrid_job(details, query):
                        skipped_workplace_type.append(_job_label(details.company, details.title))
                        query_outcome.skipped_workplace_type += 1
                        continue
                    if _is_disallowed_experience_level(details):
                        skipped_experience_level.append(
                            _job_label(details.company, details.title)
                        )
                        query_outcome.skipped_experience_level += 1
                        continue
                    candidates.append(details)
                    query_outcome.fresh_jobs_accepted += 1
                    accepted_job_outcomes[details.job_id] = query_outcome
                if (
                    len(candidates) >= max_jobs
                    and search_memory.total_searches >= min_searches_before_stop
                ):
                    break

            if not new_query_found or search_memory.total_searches == 0:
                break

        artifacts_by_job_id: dict[str, TailoredResumeArtifact] = {}
        failed_cover_letter_job_ids: set[str] = set()
        for index, job in enumerate(candidates, start=1):
            _emit_job_progress(
                progress_callback,
                action="Working on job",
                job=job,
                index=index,
                total=len(candidates),
            )
            try:
                artifact = await self._generate_and_store_application_artifacts(
                    source_resume=source_resume,
                    current_job_description=current_job_description,
                    job=job,
                    output_dir=output_dir,
                    tracking_path=tracking_path,
                    application_database_path=application_database_path,
                    append_tracking=job.job_id not in existing_application_job_ids,
                    stored_job_description=job.description,
                    artifact_mode=artifact_mode,
                    progress_callback=progress_callback,
                )
                if artifact.resume_path and artifact.artifact_kind == "resume":
                    existing_resume_job_ids.add(job.job_id)
                if artifact.cover_letter_path:
                    existing_cover_letter_job_ids.add(job.job_id)
                existing_application_job_ids.add(job.job_id)
                if outcome := accepted_job_outcomes.get(job.job_id):
                    _update_query_outcome_from_artifact(
                        outcome=outcome,
                        artifact=artifact,
                        job=job,
                    )
            except LinkedInCareerMcpError as exc:
                errors.append(f"{job.job_id}: {exc}")
                _emit_progress(
                    progress_callback,
                    f"Failed to generate artifacts for job id '{job.job_id}': {exc}",
                )
                if artifact_mode in {"all", "cover-letters-only"}:
                    failed_cover_letter_job_ids.add(job.job_id)
                continue
            artifacts_by_job_id[job.job_id] = artifact

        cover_letter_retries = await self._retry_missing_cover_letters(
            candidates=[
                _RegenerationCandidate(job=job, stored_job_description=job.description)
                for job in candidates
            ],
            source_resume=source_resume,
            current_job_description=current_job_description,
            output_dir=output_dir,
            tracking_path=tracking_path,
            application_database_path=application_database_path,
            artifacts_by_job_id=artifacts_by_job_id,
            errors=errors,
            artifact_mode=artifact_mode,
            cover_letter_retry_attempts=cover_letter_retry_attempts,
            force_retry_job_ids=failed_cover_letter_job_ids,
            progress_callback=progress_callback,
        )
        artifacts = list(artifacts_by_job_id.values())
        artifact_audit = audit_application_artifacts(application_database_path)
        _record_query_run_outcomes(
            database_path=application_database_path,
            outcomes=query_run_outcomes,
            artifact_mode=artifact_mode,
        )

        return MatchingJobsWorkflowResult(
            profile_files=[str(document.path) for document in profile_documents],
            search_queries=all_search_queries,
            jobs_found=len(candidates),
            resumes_created=sum(1 for artifact in artifacts if artifact.artifact_kind == "resume"),
            cover_letters_created=sum(1 for artifact in artifacts if artifact.cover_letter_path),
            recommendations_created=sum(
                1 for artifact in artifacts if artifact.artifact_kind == "recommendations"
            ),
            tracking_spreadsheet=str(tracking_path),
            skipped_blacklisted=skipped_blacklisted,
            skipped_existing=skipped_existing,
            skipped_workplace_type=skipped_workplace_type,
            skipped_experience_level=skipped_experience_level,
            errors=errors,
            artifacts=artifacts,
            artifact_audit=artifact_audit,
            cover_letter_retries=cover_letter_retries,
        )

    async def regenerate_resumes(
        self,
        *,
        profile_dir: Path = DEFAULT_PROFILE_DIR,
        output_dir: Path = DEFAULT_OUTPUT_DIR,
        source_resume_name: str = DEFAULT_SOURCE_RESUME,
        current_job_description_name: str = DEFAULT_CURRENT_JOB_DESCRIPTION,
        job_ids: list[str] | None = None,
        linkedin_delay_seconds: float = 2.0,
        artifact_mode: ArtifactMode = "resumes-only",
        cover_letter_retry_attempts: int = 0,
        progress_callback: ProgressCallback | None = None,
    ) -> MatchingJobsWorkflowResult:
        profile_documents = load_profile_documents(profile_dir)
        source_resume = _find_source_resume(profile_documents, source_resume_name)
        current_job_description = _find_source_resume(
            profile_documents,
            current_job_description_name,
        )
        tracking_path = output_dir / TRACKING_WORKBOOK
        application_database_path = output_dir / APPLICATION_DATABASE
        if tracking_path.exists():
            import_output_artifacts(output_dir=output_dir, database_path=application_database_path)

        normalized_job_ids = _normalize_regeneration_job_ids(job_ids)
        records = fetch_application_job_records(
            application_database_path,
            job_ids=normalized_job_ids,
        )
        errors: list[str] = []
        if normalized_job_ids is not None:
            found_job_ids = {record.job_id for record in records}
            for job_id in normalized_job_ids:
                if job_id not in found_job_ids:
                    errors.append(f"{job_id}: not found in application database.")
        if not records and normalized_job_ids is None:
            errors.append(f"No existing jobs were found in {application_database_path}.")

        candidates: list[_RegenerationCandidate] = []
        linkedin_fetch_count = 0
        for record in records:
            if _record_has_usable_job_description(record):
                candidates.append(_regeneration_candidate_from_record(record))
                continue
            try:
                if linkedin_fetch_count > 0 and linkedin_delay_seconds > 0:
                    await asyncio.sleep(linkedin_delay_seconds)
                details = await self._service.get_details(record.linkedin_url or record.job_id)
                linkedin_fetch_count += 1
            except LinkedInCareerMcpError as exc:
                errors.append(f"{record.job_id}: {exc}")
                _emit_progress(
                    progress_callback,
                    f"Failed to fetch LinkedIn details for job id '{record.job_id}': {exc}",
                )
                continue
            candidates.append(
                _RegenerationCandidate(
                    job=_merge_record_with_fetched_details(record, details),
                    stored_job_description=details.description,
                )
            )

        artifacts_by_job_id: dict[str, TailoredResumeArtifact] = {}
        failed_cover_letter_job_ids: set[str] = set()
        for index, candidate in enumerate(candidates, start=1):
            _emit_job_progress(
                progress_callback,
                action="Working on job",
                job=candidate.job,
                index=index,
                total=len(candidates),
            )
            try:
                artifact = await self._generate_and_store_application_artifacts(
                    source_resume=source_resume,
                    current_job_description=current_job_description,
                    job=candidate.job,
                    output_dir=output_dir,
                    tracking_path=tracking_path,
                    application_database_path=application_database_path,
                    append_tracking=False,
                    stored_job_description=candidate.stored_job_description,
                    artifact_mode=artifact_mode,
                    progress_callback=progress_callback,
                )
            except LinkedInCareerMcpError as exc:
                errors.append(f"{candidate.job.job_id}: {exc}")
                _emit_progress(
                    progress_callback,
                    f"Failed to generate artifacts for job id '{candidate.job.job_id}': {exc}",
                )
                if artifact_mode in {"all", "cover-letters-only"}:
                    failed_cover_letter_job_ids.add(candidate.job.job_id)
                continue
            artifacts_by_job_id[candidate.job.job_id] = artifact

        cover_letter_retries = await self._retry_missing_cover_letters(
            candidates=candidates,
            source_resume=source_resume,
            current_job_description=current_job_description,
            output_dir=output_dir,
            tracking_path=tracking_path,
            application_database_path=application_database_path,
            artifacts_by_job_id=artifacts_by_job_id,
            errors=errors,
            artifact_mode=artifact_mode,
            cover_letter_retry_attempts=cover_letter_retry_attempts,
            force_retry_job_ids=failed_cover_letter_job_ids,
            progress_callback=progress_callback,
        )
        artifacts = list(artifacts_by_job_id.values())
        artifact_audit = audit_application_artifacts(application_database_path)

        return MatchingJobsWorkflowResult(
            profile_files=[str(document.path) for document in profile_documents],
            search_queries=[],
            jobs_found=len(candidates),
            resumes_created=sum(1 for artifact in artifacts if artifact.artifact_kind == "resume"),
            cover_letters_created=sum(1 for artifact in artifacts if artifact.cover_letter_path),
            recommendations_created=sum(
                1 for artifact in artifacts if artifact.artifact_kind == "recommendations"
            ),
            tracking_spreadsheet=str(tracking_path),
            skipped_blacklisted=[],
            skipped_existing=[],
            skipped_workplace_type=[],
            skipped_experience_level=[],
            errors=errors,
            artifacts=artifacts,
            artifact_audit=artifact_audit,
            cover_letter_retries=cover_letter_retries,
        )

    async def _retry_missing_cover_letters(
        self,
        *,
        candidates: list[_RegenerationCandidate],
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        output_dir: Path,
        tracking_path: Path,
        application_database_path: Path,
        artifacts_by_job_id: dict[str, TailoredResumeArtifact],
        errors: list[str],
        artifact_mode: ArtifactMode,
        cover_letter_retry_attempts: int,
        force_retry_job_ids: set[str] | None,
        progress_callback: ProgressCallback | None,
    ) -> list[CoverLetterRetryRecord]:
        if artifact_mode not in {"all", "cover-letters-only"}:
            return []
        retry_attempts = max(0, cover_letter_retry_attempts)
        if retry_attempts == 0 or not candidates:
            return []

        retry_records: list[CoverLetterRetryRecord] = []
        remaining_force_retry_job_ids = set(force_retry_job_ids or set())
        for attempt in range(1, retry_attempts + 1):
            existing_cover_letter_job_ids = fetch_existing_cover_letter_job_ids(
                application_database_path
            )
            missing_candidates = [
                candidate
                for candidate in candidates
                if candidate.job.job_id not in existing_cover_letter_job_ids
                or candidate.job.job_id in remaining_force_retry_job_ids
            ]
            if not missing_candidates:
                break

            for index, candidate in enumerate(missing_candidates, start=1):
                _emit_job_progress(
                    progress_callback,
                    action="Retrying cover letter",
                    job=candidate.job,
                    index=index,
                    total=len(missing_candidates),
                    detail=f"attempt {attempt}/{retry_attempts}",
                )
                try:
                    artifact = await self._generate_and_store_application_artifacts(
                        source_resume=source_resume,
                        current_job_description=current_job_description,
                        job=candidate.job,
                        output_dir=output_dir,
                        tracking_path=tracking_path,
                        application_database_path=application_database_path,
                        append_tracking=False,
                        stored_job_description=candidate.stored_job_description,
                        artifact_mode="cover-letters-only",
                        progress_callback=progress_callback,
                    )
                except LinkedInCareerMcpError as exc:
                    error = f"{candidate.job.job_id}: cover letter retry {attempt}: {exc}"
                    errors.append(error)
                    retry_records.append(
                        CoverLetterRetryRecord(
                            attempt=attempt,
                            job_id=candidate.job.job_id,
                            company=candidate.job.company,
                            title=candidate.job.title,
                            status="failed",
                            error=str(exc),
                        )
                    )
                    continue

                artifacts_by_job_id[candidate.job.job_id] = _merge_artifacts(
                    artifacts_by_job_id.get(candidate.job.job_id),
                    artifact,
                )
                remaining_force_retry_job_ids.discard(candidate.job.job_id)
                retry_records.append(
                    CoverLetterRetryRecord(
                        attempt=attempt,
                        job_id=candidate.job.job_id,
                        company=candidate.job.company,
                        title=candidate.job.title,
                        status="created",
                        cover_letter_path=artifact.cover_letter_path,
                    )
                )

        return retry_records

    async def _generate_and_store_application_artifacts(
        self,
        *,
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        job: JobDetails,
        output_dir: Path,
        tracking_path: Path,
        application_database_path: Path,
        append_tracking: bool,
        stored_job_description: str | None = None,
        artifact_mode: ArtifactMode,
        progress_callback: ProgressCallback | None = None,
    ) -> TailoredResumeArtifact:
        artifact_kind: Literal["resume", "recommendations", "cover_letter"] = "cover_letter"
        resume_path: Path | None = None
        recommendations_path: Path | None = None
        cover_letter_path: Path | None = None

        if artifact_mode in {"all", "resumes-only"}:
            _emit_progress(progress_callback, f"Generating resume for job id '{job.job_id}'.")
            resume_draft = await self._generate_resume_draft(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
            )
            resume_text = resume_draft.text
            artifact_kind = "resume"
            if _looks_like_recommendations(resume_text):
                artifact_kind = "recommendations"
                recommendations_text = await self._generate_recommendations_text(
                    source_resume=source_resume,
                    current_job_description=current_job_description,
                    job=job,
                    draft_text=resume_text,
                )
                resume_path = write_resume_recommendations_pdf(
                    recommendations_text=recommendations_text,
                    output_dir=output_dir,
                    job=job,
                )
                _emit_progress(
                    progress_callback,
                    f"Wrote resume recommendations for job id '{job.job_id}': {resume_path}",
                )
                recommendations_path = resume_path
            else:
                resume_path = write_resume_pdf(
                    resume_text=resume_text,
                    output_dir=output_dir,
                    job=job,
                )
                _emit_progress(
                    progress_callback,
                    f"Wrote resume PDF for job id '{job.job_id}': {resume_path}",
                )
                repaired_draft, repaired_resume_path, repair_score = (
                    await self._repair_resume_draft_from_ats(
                        draft=resume_draft,
                        source_resume=source_resume,
                        job=job,
                        output_dir=output_dir,
                        resume_path=resume_path,
                        progress_callback=progress_callback,
                    )
                )
                resume_text = repaired_draft.text
                resume_path = repaired_resume_path
                repair_target_met = (
                    repair_score is not None
                    and repair_score.overall_score >= ATS_REPAIR_TARGET_SCORE
                )
                if repair_target_met:
                    _emit_progress(
                        progress_callback,
                        "ATS repair target met "
                        f"({repair_score.overall_score}/100) for job id '{job.job_id}'.",
                    )

        if artifact_mode in {"all", "cover-letters-only"}:
            _emit_progress(
                progress_callback,
                f"Generating cover letter for job id '{job.job_id}'.",
            )
            cover_letter_text = await self._generate_cover_letter_text(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
            )
            cover_letter_path = write_cover_letter_pdf(
                cover_letter_text=cover_letter_text,
                output_dir=output_dir,
                job=job,
            )
            _emit_progress(
                progress_callback,
                f"Wrote cover letter PDF for job id '{job.job_id}': {cover_letter_path}",
            )

        if append_tracking:
            append_tracking_row(
                tracking_path=tracking_path,
                job=job,
                resume_path=resume_path,
                cover_letter_path=cover_letter_path,
            )
        prompt_job_description = _job_description_context(job)
        upsert_application_artifact(
            database_path=application_database_path,
            job_id=job.job_id,
            company=job.company or "",
            job_title=job.title,
            linkedin_url=str(job.job_url) if job.job_url else "",
            resume_path=resume_path,
            cover_letter_path=cover_letter_path,
            job_description=stored_job_description,
            prompt_job_description=_usable_job_description(prompt_job_description),
            date_posted=_job_date_posted(job),
            experience_level=job.seniority_level,
        )
        _emit_progress(
            progress_callback,
            f"Stored artifacts in database for job id '{job.job_id}'.",
        )
        return TailoredResumeArtifact(
            job_id=job.job_id,
            company=job.company,
            title=job.title,
            linkedin_url=str(job.job_url) if job.job_url else None,
            resume_path=str(resume_path) if resume_path else "",
            cover_letter_path=str(cover_letter_path) if cover_letter_path else None,
            artifact_kind=artifact_kind,
            recommendations_path=str(recommendations_path) if recommendations_path else None,
        )

    async def _generate_search_queries(
        self,
        *,
        profile_context: str,
        location: str,
        date_posted: DatePosted,
        limit_per_query: int,
        max_queries: int,
        query_history: list[StoredQueryOutcome],
        search_memory: _SearchMemory | None = None,
    ) -> list[ScoredQuery]:
        try:
            plan = await self._planner_llm.generate_json(
                _search_query_prompt(
                    profile_context=_limit_context(profile_context, max_chars=8_000),
                    location=location,
                    date_posted=date_posted,
                    limit_per_query=limit_per_query,
                    max_queries=max_queries,
                    search_memory=search_memory,
                )
            )
        except WorkflowError:
            # Search planning is best-effort; history and local supplements can still
            # produce useful LinkedIn queries during a temporary LLM outage.
            plan = {}
        raw_queries = plan.get("queries")
        if raw_queries is None and "keywords" in plan:
            raw_queries = [plan]
        if not isinstance(raw_queries, list):
            raw_queries = []

        base_queries: list[JobSearchQuery] = []
        for value in raw_queries:
            if not isinstance(value, dict):
                continue
            try:
                base_queries.append(
                    _coerce_search_query(
                        value,
                        location=location,
                        date_posted=date_posted,
                        limit_per_query=limit_per_query,
                    )
                )
            except WorkflowError:
                continue

        supplemented_queries = _supplement_search_queries(
            [
                *base_queries,
                *historical_query_candidates(
                    query_history,
                    location=location,
                    date_posted=date_posted,
                    limit_per_query=limit_per_query,
                ),
            ],
            location=location,
            date_posted=date_posted,
            limit_per_query=limit_per_query,
        )
        if not supplemented_queries:
            raise WorkflowError("The LLM did not return usable LinkedIn search queries.")
        candidate_pool = _expand_remote_and_hybrid_queries(
            supplemented_queries,
            max_queries=max(max_queries * 3, max_queries),
        )
        return rank_search_queries(
            candidate_pool,
            profile_context=profile_context,
            history=query_history,
            max_queries=max_queries,
        )

    async def _generate_resume_text(
        self,
        *,
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        job: JobDetails,
    ) -> str:
        return (
            await self._generate_resume_draft(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
            )
        ).text

    async def _generate_resume_draft(
        self,
        *,
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        job: JobDetails,
    ) -> _ResumeDraft:
        if source_resume is None:
            raise WorkflowError(f"Source resume file was not found: {DEFAULT_SOURCE_RESUME}")
        tailored_scjdir = await self._ollama.generate_text(
            _scjdir_prompt(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
            )
        )
        if not tailored_scjdir:
            raise WorkflowError(f"The LLM returned an empty SCJDiR rewrite for job {job.job_id}.")
        sections_plan = await self._ollama.generate_json(
            _resume_sections_prompt(
                source_resume=source_resume,
                current_job_description=current_job_description,
                tailored_scjdir=tailored_scjdir,
                job=job,
            )
        )
        text = _render_resume_template(
            tailored_scjdir=tailored_scjdir,
            sections_plan=sections_plan,
        )
        return _ResumeDraft(
            tailored_scjdir=tailored_scjdir,
            sections_plan=sections_plan if isinstance(sections_plan, Mapping) else {},
            text=text,
        )

    async def _repair_resume_draft_from_ats(
        self,
        *,
        draft: _ResumeDraft,
        source_resume: ProfileDocument | None,
        job: JobDetails,
        output_dir: Path,
        resume_path: Path,
        progress_callback: ProgressCallback | None = None,
    ) -> tuple[_ResumeDraft, Path, AtsProxyScore | None]:
        if source_resume is None or not resume_path.is_file():
            return draft, resume_path, None

        job_description = _job_description_context(job)
        if not job_description.strip():
            return draft, resume_path, None

        current_draft = draft
        current_path = resume_path
        best_draft = draft
        best_path = resume_path
        best_score = calculate_ats_proxy_score(
            resume_pdf=resume_path.read_bytes(),
            job_description=job_description,
        )

        for attempt in range(1, ATS_REPAIR_MAX_ATTEMPTS + 1):
            if best_score.overall_score >= ATS_REPAIR_TARGET_SCORE:
                break
            missing_terms = best_score.missing_high_value_terms[
                :ATS_REPAIR_MAX_MISSING_TERMS
            ]
            source_evidence = _source_resume_evidence_for_missing_terms(
                source_resume_text=source_resume.text,
                missing_terms=missing_terms,
            )
            if not source_evidence:
                _emit_progress(
                    progress_callback,
                    "ATS repair skipped; no source-resume evidence for missing terms "
                    f"on job id '{job.job_id}'.",
                )
                break

            _emit_progress(
                progress_callback,
                "ATS repair attempt "
                f"{attempt}/{ATS_REPAIR_MAX_ATTEMPTS} for job id '{job.job_id}' "
                f"(score {best_score.overall_score}/100).",
            )
            repair_plan = await self._ollama.generate_json(
                _ats_resume_repair_prompt(
                    source_evidence=source_evidence,
                    current_resume_text=current_draft.text,
                    current_tailored_scjdir=current_draft.tailored_scjdir,
                    current_sections_plan=current_draft.sections_plan,
                    job=job,
                    score=best_score,
                )
            )
            repaired_draft = _coerce_ats_repaired_resume_draft(
                repair_plan=repair_plan,
                current_draft=current_draft,
            )
            if repaired_draft.text == current_draft.text:
                _emit_progress(
                    progress_callback,
                    f"ATS repair made no changes for job id '{job.job_id}'.",
                )
                break

            repaired_path = write_resume_pdf(
                resume_text=repaired_draft.text,
                output_dir=output_dir,
                job=job,
            )
            repaired_score = calculate_ats_proxy_score(
                resume_pdf=repaired_path.read_bytes(),
                job_description=job_description,
            )
            if repaired_score.overall_score <= best_score.overall_score:
                write_resume_pdf(resume_text=best_draft.text, output_dir=output_dir, job=job)
                _emit_progress(
                    progress_callback,
                    "ATS repair kept existing best resume "
                    f"({best_score.overall_score}/100) for job id '{job.job_id}'.",
                )
                break

            _emit_progress(
                progress_callback,
                "ATS repair improved score "
                f"{best_score.overall_score}/100 -> {repaired_score.overall_score}/100 "
                f"for job id '{job.job_id}'.",
            )
            best_draft = repaired_draft
            best_path = repaired_path
            best_score = repaired_score
            current_draft = repaired_draft
            current_path = repaired_path

        return best_draft, best_path or current_path, best_score

    async def _generate_cover_letter_text(
        self,
        *,
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        job: JobDetails,
    ) -> str:
        if source_resume is None:
            raise WorkflowError(f"Source resume file was not found: {DEFAULT_SOURCE_RESUME}")
        sections_plan = await self._ollama.generate_json(
            _cover_letter_sections_prompt(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
            )
        )
        return _render_cover_letter_template(job=job, sections_plan=sections_plan)

    async def _generate_recommendations_text(
        self,
        *,
        source_resume: ProfileDocument | None,
        current_job_description: ProfileDocument | None,
        job: JobDetails,
        draft_text: str,
    ) -> str:
        text = await self._ollama.generate_text(
            _recommendations_prompt(
                source_resume=source_resume,
                current_job_description=current_job_description,
                job=job,
                draft_text=draft_text,
            )
        )
        if not text:
            raise WorkflowError(f"The LLM returned empty recommendations for job {job.job_id}.")
        return text


def load_profile_documents(profile_dir: Path) -> list[ProfileDocument]:
    if not profile_dir.exists():
        raise WorkflowError(f"Profile directory does not exist: {profile_dir}")
    documents: list[ProfileDocument] = []
    for path in sorted(profile_dir.rglob("*")):
        if not path.is_file() or path.name.startswith("."):
            continue
        if path.suffix.lower() not in SUPPORTED_PROFILE_SUFFIXES:
            continue
        text = _read_profile_file(path)
        if text.strip():
            documents.append(ProfileDocument(path=path, text=text.strip()))
    if not documents:
        supported = ", ".join(sorted(SUPPORTED_PROFILE_SUFFIXES))
        raise WorkflowError(f"No supported profile files found in {profile_dir} ({supported}).")
    return documents



def _source_resume_evidence_for_missing_terms(
    *,
    source_resume_text: str,
    missing_terms: tuple[str, ...],
) -> dict[str, str]:
    evidence: dict[str, str] = {}
    for term in missing_terms[:ATS_REPAIR_MAX_MISSING_TERMS]:
        snippet = _source_resume_snippet_for_term(source_resume_text, term)
        if snippet is None:
            continue
        evidence[term] = snippet
        if len(evidence) >= ATS_REPAIR_MAX_SOURCE_SNIPPETS:
            break
    return evidence


def _source_resume_snippet_for_term(source_resume_text: str, term: str) -> str | None:
    term_key = _normalize_source_match_key(term)
    if not term_key or term_key in {"ai", "api", "cloud"}:
        return None
    lines = [line.strip() for line in source_resume_text.splitlines() if line.strip()]
    for index, line in enumerate(lines):
        if _source_line_matches_term(line=line, term_key=term_key):
            start = max(0, index - ATS_REPAIR_SOURCE_SNIPPET_RADIUS)
            end = min(len(lines), index + ATS_REPAIR_SOURCE_SNIPPET_RADIUS + 1)
            return " ".join(lines[start:end])[:900]
    alias_snippet = _source_resume_alias_snippet(lines=lines, term_key=term_key)
    if alias_snippet:
        return alias_snippet
    normalized_source = _normalize_source_match_key(source_resume_text)
    if term_key not in normalized_source:
        return None
    raw_index = normalized_source.find(term_key)
    window_start = max(0, raw_index - 360)
    window_end = min(len(source_resume_text), raw_index + len(term) + 360)
    return _clean_inline_text(source_resume_text[window_start:window_end])[:900]


def _source_line_matches_term(*, line: str, term_key: str) -> bool:
    line_key = _normalize_source_match_key(line)
    if term_key in line_key:
        return True
    if term_key == "restapi" and "restfulapi" in line_key:
        return True
    if term_key == "rbac" and "rolebasedaccesscontrol" in line_key:
        return True
    return term_key == "cicd" and "continuousintegration" in line_key


def _source_resume_alias_snippet(*, lines: list[str], term_key: str) -> str | None:
    aliases = SOURCE_EVIDENCE_TERM_ALIASES.get(term_key)
    if not aliases:
        return None
    snippets: list[str] = []
    matched_aliases: set[str] = set()
    for index, line in enumerate(lines):
        for alias in aliases:
            if alias in matched_aliases:
                continue
            if not _source_line_matches_alias(line=line, alias_key=alias):
                continue
            start = max(0, index - ATS_REPAIR_SOURCE_SNIPPET_RADIUS)
            end = min(len(lines), index + ATS_REPAIR_SOURCE_SNIPPET_RADIUS + 1)
            snippets.append(" ".join(lines[start:end]))
            matched_aliases.add(alias)
            break
        if len(snippets) >= 2:
            break
    if not snippets:
        return None
    return " ".join(snippets)[:900]


def _source_line_matches_alias(*, line: str, alias_key: str) -> bool:
    if alias_key == "ai":
        return bool(re.search(r"(?<![a-z0-9])ai(?:[-\s]|$)", line, flags=re.IGNORECASE))
    line_key = _normalize_source_match_key(line)
    if alias_key in line_key:
        return True
    if alias_key == "cicd" and "continuousintegration" in line_key:
        return True
    return alias_key == "developerworkflow" and "developertooling" in line_key


def _normalize_source_match_key(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.casefold())


def format_profile_context(documents: list[ProfileDocument], *, max_chars: int = 120_000) -> str:
    sections: list[str] = []
    remaining = max_chars
    for document in documents:
        if remaining <= 0:
            break
        text = document.text[:remaining]
        sections.append(f"--- {document.path.name} ---\n{text}")
        remaining -= len(text)
    return "\n\n".join(sections)


def _limit_context(text: str, *, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rsplit("\n", 1)[0]


def write_resume_pdf(*, resume_text: str, output_dir: Path, job: JobDetails) -> Path:
    company_dir = _path_part(job.company or "unknown_company")
    title_part = _path_part(job.title, lower=True)
    job_dir = _path_part(f"{job.job_id}_{job.title}", lower=True)
    resume_path = output_dir / "resumes" / company_dir / job_dir / f"mp_resume_{title_part}.pdf"
    _write_text_pdf(text=resume_text, path=resume_path)
    return resume_path


def write_resume_recommendations_pdf(
    *,
    recommendations_text: str,
    output_dir: Path,
    job: JobDetails,
) -> Path:
    company_dir = _path_part(job.company or "unknown_company")
    title_part = _path_part(job.title, lower=True)
    job_dir = _path_part(f"{job.job_id}_{job.title}", lower=True)
    recommendations_path = (
        output_dir
        / "resumes"
        / company_dir
        / job_dir
        / f"mp_resume_{title_part}-recommends.pdf"
    )
    _write_text_pdf(text=recommendations_text, path=recommendations_path)
    return recommendations_path


def write_cover_letter_pdf(*, cover_letter_text: str, output_dir: Path, job: JobDetails) -> Path:
    company_dir = _path_part(job.company or "unknown_company")
    title_part = _path_part(job.title, lower=True)
    job_dir = _path_part(f"{job.job_id}_{job.title}", lower=True)
    cover_letter_path = (
        output_dir
        / "cover_letters"
        / company_dir
        / job_dir
        / f"mp_cover_letter_{title_part}.pdf"
    )
    _write_cover_letter_text_pdf(text=cover_letter_text, path=cover_letter_path)
    return cover_letter_path


def _write_cover_letter_text_pdf(*, text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CoverLetterBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14.5,
        spaceAfter=10,
        textColor=RESUME_BODY_COLOR,
    )
    signature = ParagraphStyle(
        "CoverLetterSignature",
        parent=body,
        spaceBefore=2,
        spaceAfter=2,
    )

    story: list[Any] = [_cover_letter_top_bar(), Spacer(1, 34)]
    content_added = False
    for paragraph in _clean_cover_letter_text(text).split("\n\n"):
        line = paragraph.strip()
        if not line:
            continue
        paragraph_lines = [part.strip() for part in line.splitlines() if part.strip()]
        style = (
            signature
            if paragraph_lines
            and all(
                part in {"Sincerely,", "Maxim Perkhounkov", LINKEDIN_PROFILE_MARKDOWN}
                for part in paragraph_lines
            )
            else body
        )
        story.append(Paragraph(_cover_letter_markup(line), style))
        content_added = True

    if not content_added:
        story.append(Paragraph("Cover letter content was empty.", body))

    document = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=72,
        leftMargin=72,
        topMargin=54,
        bottomMargin=72,
    )
    document.build(story)


def _write_text_pdf(*, text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ResumeBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.8,
        leading=10.6,
        spaceAfter=2,
        textColor=RESUME_BODY_COLOR,
    )
    bullet = ParagraphStyle(
        "ResumeBullet",
        parent=body,
        leftIndent=0.32 * inch,
        firstLineIndent=-0.18 * inch,
        bulletIndent=0.03 * inch,
        spaceAfter=1,
    )
    nested_bullet = ParagraphStyle(
        "ResumeNestedBullet",
        parent=body,
        leftIndent=0.72 * inch,
        firstLineIndent=-0.14 * inch,
        bulletIndent=0.46 * inch,
        spaceAfter=1,
    )
    heading = ParagraphStyle(
        "ResumeHeading",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=12.2,
        leading=14.2,
        spaceBefore=9,
        spaceAfter=6,
        textColor=EMERALD_DARK,
    )
    employer_style = ParagraphStyle(
        "ResumeEmployer",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=9.5,
        leading=11.4,
        spaceBefore=9,
        spaceAfter=5,
        textColor=RESUME_BODY_COLOR,
    )
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=8.6,
        leading=10.3,
        spaceAfter=8,
        textColor=RESUME_MUTED_COLOR,
    )
    name_style = ParagraphStyle(
        "ResumeName",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=13.2,
        leading=15.4,
        alignment=0,
        spaceAfter=5,
        textColor=EMERALD_DARK,
    )
    contact_style = ParagraphStyle(
        "ResumeContact",
        parent=body,
        fontName="Helvetica-Bold",
        fontSize=7.2,
        leading=8.6,
        alignment=0,
        textColor=RESUME_MUTED_COLOR,
        spaceAfter=8,
    )
    note_style = ParagraphStyle(
        "ResumeNote",
        parent=body,
        fontName="Helvetica-Oblique",
        fontSize=8,
        leading=9.6,
        textColor=RESUME_MUTED_COLOR,
    )

    story: list[Any] = []
    story.append(_resume_rule(top=True))
    story.append(Spacer(1, 12))

    current_section: str | None = None
    previous_line_blank = False
    for line_number, raw_line in enumerate(_clean_resume_text(text).splitlines()):
        raw_line = _strip_markdown_emphasis(raw_line.rstrip())
        line = raw_line.strip()
        is_nested_bullet = raw_line.startswith("  - ")
        if not line:
            if not previous_line_blank:
                story.append(Spacer(1, 5))
            previous_line_blank = True
            continue
        previous_line_blank = False

        extra_section_heading = _looks_like_extra_resume_section_heading(
            line,
            current_section=current_section,
        )
        if line in RESUME_SECTION_HEADINGS or extra_section_heading:
            if (
                line in {"Professional Experience", "Education & Certifications"}
                or extra_section_heading
            ):
                story.append(Spacer(1, 5))
                story.append(_resume_rule())
                story.append(Spacer(1, 8))
            current_section = line
            story.append(Paragraph(_paragraph_markup(line), heading))
        elif line_number == 0:
            story.append(Paragraph(_paragraph_markup(line), name_style))
        elif current_section is None:
            story.append(Paragraph(_paragraph_markup(line), contact_style))
        elif line.startswith("Note:"):
            story.append(Paragraph(_paragraph_markup(line), note_style))
        elif current_section == "Professional Experience" and _looks_like_employer_line(line):
            story.append(Paragraph(_paragraph_markup(line), employer_style))
        elif current_section == "Professional Experience" and _looks_like_title_line(line):
            story.append(Paragraph(_title_markup(line), title_style))
        elif is_nested_bullet:
            story.append(
                Paragraph(
                    _nested_bullet_markup(raw_line[4:].strip()),
                    nested_bullet,
                    bulletText="o",
                )
            )
        elif line.startswith("- "):
            story.append(Paragraph(_bullet_markup(line[2:]), bullet, bulletText="\u2022"))
        elif _looks_like_heading(line):
            story.append(Paragraph(_paragraph_markup(line), heading))
        else:
            story.append(Paragraph(_paragraph_markup(line), body))

    if not story:
        story.append(Paragraph("Resume content was empty.", body))

    document = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        rightMargin=54,
        leftMargin=54,
        topMargin=58,
        bottomMargin=52,
    )
    document.build(story)


def append_tracking_row(
    *,
    tracking_path: Path,
    job: JobDetails,
    resume_path: Path | None,
    cover_letter_path: Path | None = None,
) -> None:
    tracking_path.parent.mkdir(parents=True, exist_ok=True)
    workbook: Workbook
    if tracking_path.exists():
        workbook = load_workbook(tracking_path)
        sheet = workbook.active
        _ensure_tracking_headers(sheet)
    else:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Applications"
        sheet.append(TRACKING_HEADERS)
        sheet.freeze_panes = "A2"
        _size_tracking_columns(sheet)

    row = sheet.max_row + 1
    job_url = str(job.job_url) if job.job_url else ""
    relative_resume = str(resume_path) if resume_path else ""
    relative_cover_letter = str(cover_letter_path) if cover_letter_path else ""
    values = [
        job.job_id,
        job.company or "",
        job.title,
        job_url,
        relative_resume,
        relative_cover_letter,
        "No",
        "",
    ]
    for column, value in enumerate(values, start=1):
        sheet.cell(row=row, column=column, value=value)

    if job_url:
        link_cell = sheet.cell(row=row, column=4)
        link_cell.hyperlink = job_url
        link_cell.style = "Hyperlink"

    if resume_path is not None:
        resume_cell = sheet.cell(row=row, column=5)
        resume_cell.hyperlink = resume_path.resolve().as_uri()
        resume_cell.style = "Hyperlink"

    if cover_letter_path is not None:
        cover_letter_cell = sheet.cell(row=row, column=6)
        cover_letter_cell.hyperlink = cover_letter_path.resolve().as_uri()
        cover_letter_cell.style = "Hyperlink"

    workbook.save(tracking_path)


def _read_profile_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in SUPPORTED_TEXT_SUFFIXES:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document

        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return ""


def _find_source_resume(
    documents: list[ProfileDocument],
    source_resume_name: str,
) -> ProfileDocument | None:
    for document in documents:
        if document.path.name == source_resume_name:
            return document
    return None


def _normalize_regeneration_job_ids(job_ids: list[str] | None) -> list[str] | None:
    if job_ids is None:
        return None
    normalized = _dedupe_preserve_order(_split_regeneration_job_ids(job_ids))
    if not normalized:
        return None
    all_markers = [job_id for job_id in normalized if job_id.casefold() == "all"]
    if all_markers and len(normalized) > 1:
        raise WorkflowError("Use either 'all' or explicit LinkedIn job IDs, not both.")
    if all_markers:
        return None
    return normalized


def _split_regeneration_job_ids(job_ids: list[str]) -> list[str]:
    values: list[str] = []
    for raw_value in job_ids:
        for job_id in str(raw_value).split(","):
            value = job_id.strip()
            if value:
                values.append(value)
    return values


def _regeneration_candidate_from_record(record: ApplicationJobRecord) -> _RegenerationCandidate:
    job_description = _record_job_description(record)
    return _RegenerationCandidate(
        job=JobDetails(
            job_id=record.job_id,
            title=record.job_title or "Unknown title",
            company=record.company or None,
            job_url=record.linkedin_url or None,
            listed_at=record.date_posted,
            description=job_description,
            seniority_level=record.experience_level,
        ),
        stored_job_description=_usable_job_description(record.job_description),
    )


def _record_has_usable_job_description(record: ApplicationJobRecord) -> bool:
    return _record_job_description(record) is not None


def _record_job_description(record: ApplicationJobRecord) -> str | None:
    return (
        _usable_job_description(record.job_description)
        or _usable_job_description(record.prompt_job_description)
    )


def _usable_job_description(value: str | None) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    if _is_placeholder_job_description(text):
        return None
    return text


def _is_placeholder_job_description(value: str) -> bool:
    normalized = re.sub(r"\s+", " ", value).strip().casefold().rstrip(".")
    placeholder = NO_PUBLIC_JOB_DESCRIPTION.casefold().rstrip(".")
    return normalized == placeholder


def _update_query_outcome_from_artifact(
    *,
    outcome: _QueryRunOutcome,
    artifact: TailoredResumeArtifact,
    job: JobDetails,
) -> None:
    if artifact.artifact_kind == "resume":
        outcome.resumes_generated += 1
    if not artifact.resume_path or artifact.artifact_kind != "resume":
        return
    resume_path = Path(artifact.resume_path)
    if not resume_path.is_file():
        return
    score = calculate_ats_proxy_score(
        resume_pdf=resume_path.read_bytes(),
        job_description=_job_description_context(job),
    )
    outcome.ats_scores.append(score.overall_score)


def _record_query_run_outcomes(
    *,
    database_path: Path,
    outcomes: list[_QueryRunOutcome],
    artifact_mode: ArtifactMode,
) -> None:
    for outcome in outcomes:
        record_query_outcome(
            database_path,
            outcome.to_query_outcome(artifact_mode=artifact_mode),
        )


def _merge_record_with_fetched_details(
    record: ApplicationJobRecord,
    details: JobDetails,
) -> JobDetails:
    record_description = _record_job_description(record)
    return JobDetails(
        job_id=record.job_id,
        title=details.title if details.title != "Unknown title" else record.job_title,
        company=details.company or record.company or None,
        location=details.location,
        listed_at=details.listed_at,
        posted_text=details.posted_text,
        job_url=details.job_url or record.linkedin_url or None,
        company_url=details.company_url,
        workplace_type=details.workplace_type,
        source=details.source,
        description=details.description or record_description,
        seniority_level=details.seniority_level or record.experience_level,
        employment_type=details.employment_type,
        job_function=details.job_function,
        industries=details.industries,
    )


def _merge_posting_with_fetched_details(
    posting: JobPosting,
    details: JobDetails,
) -> JobDetails:
    return details.model_copy(
        update={
            "title": details.title if details.title != "Unknown title" else posting.title,
            "company": details.company or posting.company,
            "location": details.location or posting.location,
            "listed_at": details.listed_at or posting.listed_at,
            "posted_text": details.posted_text or posting.posted_text,
            "job_url": details.job_url or posting.job_url,
            "company_url": details.company_url or posting.company_url,
            "workplace_type": details.workplace_type or posting.workplace_type,
            "source": details.source or posting.source,
        },
    )


def _is_remote_or_hybrid_job(job: JobDetails, query: JobSearchQuery) -> bool:
    workplace_context = " ".join(
        value for value in (job.workplace_type, job.location) if value
    ).casefold()
    compact_context = re.sub(r"[\s-]+", "", workplace_context)
    if "onsite" in compact_context:
        return False
    if "remote" in workplace_context or "hybrid" in workplace_context:
        return True
    return query.workplace_type in {"remote", "hybrid"}


def _is_disallowed_experience_level(job: JobDetails) -> bool:
    normalized = _normalize_experience_level(job.seniority_level)
    return normalized in DISALLOWED_EXPERIENCE_LEVELS


def _normalize_experience_level(value: Any) -> str:
    normalized = re.sub(r"[^a-z0-9]+", "_", str(value or "").casefold()).strip("_")
    if normalized in {"intern", "internship"}:
        return "internship"
    if normalized in {"entry", "entry_level", "entrylevel"}:
        return "entry_level"
    if normalized.startswith("entry_level"):
        return "entry_level"
    return normalized


def _job_date_posted(job: JobDetails) -> str | None:
    return _posted_date_value(job.listed_at) or _posted_date_value(job.posted_text)


def _posted_date_value(value: Any) -> str | None:
    text = str(value or "").strip()
    if not text:
        return None
    date_match = re.match(r"^\d{4}-\d{2}-\d{2}", text)
    if date_match:
        return date_match.group(0)
    return text


class _JodChunkClassifier:
    def __init__(
        self,
        *,
        class_counts: Counter[str],
        token_counts: dict[str, Counter[str]],
    ) -> None:
        self.class_counts = class_counts
        self.token_counts = token_counts
        self.vocabulary = {
            token for counts in token_counts.values() for token in counts
        }
        self.token_totals = {
            label: sum(counts.values()) for label, counts in token_counts.items()
        }

    @classmethod
    def train(cls, examples: tuple[tuple[str, str], ...]) -> _JodChunkClassifier:
        class_counts: Counter[str] = Counter()
        token_counts: dict[str, Counter[str]] = {"keep": Counter(), "drop": Counter()}
        for label, text in examples:
            class_counts[label] += 1
            token_counts.setdefault(label, Counter()).update(_jod_chunk_feature_tokens(text))
        return cls(class_counts=class_counts, token_counts=token_counts)

    def keep_log_odds(self, text: str) -> float:
        tokens = _jod_chunk_feature_tokens(text)
        keep_log = self._label_log_probability(label="keep", tokens=tokens)
        drop_log = self._label_log_probability(label="drop", tokens=tokens)
        return keep_log - drop_log

    def _label_log_probability(self, *, label: str, tokens: list[str]) -> float:
        label_count = self.class_counts.get(label, 0)
        total_classes = sum(self.class_counts.values())
        class_prior = (label_count + 1) / (total_classes + max(len(self.class_counts), 1))
        token_counts = self.token_counts.get(label, Counter())
        denominator = self.token_totals.get(label, 0) + max(len(self.vocabulary), 1)
        score = math.log(class_prior)
        for token in tokens:
            score += math.log((token_counts.get(token, 0) + 1) / denominator)
        return score


def _job_description_context(job: JobDetails) -> str:
    description = _clean_job_description_for_prompt(
        job.description or NO_PUBLIC_JOB_DESCRIPTION
    )
    return _limit_context(description, max_chars=JOB_DESCRIPTION_PROMPT_MAX_CHARS)


def _clean_job_description_for_prompt(description: str) -> str:
    original = _normalize_job_description_text(description)
    if not original:
        return NO_PUBLIC_JOB_DESCRIPTION

    cleaned = _trim_low_signal_preamble(original)
    cleaned = _trim_trailing_boilerplate(cleaned)
    cleaned = _select_relevant_job_description_chunks(cleaned)
    return cleaned.strip() or original


def _trim_low_signal_preamble(description: str) -> str:
    role_start = _first_heading_match(description, ROLE_RELEVANT_START_HEADINGS)
    if role_start is None or role_start.start() == 0:
        return description

    prefix = description[: role_start.start()]
    has_low_signal_prefix = _first_heading_match(prefix, LOW_SIGNAL_PREAMBLE_HEADINGS) is not None
    if has_low_signal_prefix:
        return description[role_start.start() :].lstrip(" :-\n")
    if len(prefix) > 3_000 and _role_relevant_marker_count(prefix) < 3:
        return description[role_start.start() :].lstrip(" :-\n")
    return description


def _trim_trailing_boilerplate(description: str) -> str:
    last_role_start = max(
        (match.start() for match in _heading_matches(description, ROLE_RELEVANT_START_HEADINGS)),
        default=-1,
    )
    boilerplate = _first_heading_match(
        description,
        TRAILING_BOILERPLATE_HEADINGS,
        start=max(1, min(len(description), 300)),
        skip_before=last_role_start,
        strict_single_word=True,
    )
    if boilerplate is None:
        return description
    return description[: boilerplate.start()].rstrip(" :-\n")


def _select_relevant_job_description_chunks(description: str) -> str:
    if not _contains_any_casefolded(description, JOD_HARD_DROP_MARKERS):
        return description

    chunks = _job_description_chunks(description)
    if len(chunks) <= 1:
        return description

    kept_chunks: list[str] = []
    dropped_count = 0
    seen_role_chunk = False
    for chunk in chunks:
        contains_role_heading = _contains_role_relevant_heading(chunk)
        if _keep_job_description_chunk(
            chunk,
            preserve_company_context=not seen_role_chunk,
        ):
            kept_chunks.append(chunk)
        else:
            dropped_count += 1
        seen_role_chunk = seen_role_chunk or contains_role_heading

    selected = "\n".join(kept_chunks).strip()
    if dropped_count == 0 or not selected:
        return description
    if len(selected) < min(900, len(description) * 0.35):
        return description
    return selected


def _job_description_chunks(description: str) -> list[str]:
    boundary_headings = (
        *ROLE_RELEVANT_START_HEADINGS,
        *LOW_SIGNAL_PREAMBLE_HEADINGS,
        *TRAILING_BOILERPLATE_HEADINGS,
    )
    boundaries = {0, len(description)}
    for match in _heading_matches(description, boundary_headings):
        boundaries.add(match.start())

    chunks: list[str] = []
    sorted_boundaries = sorted(boundaries)
    for index, start in enumerate(sorted_boundaries[:-1]):
        end = sorted_boundaries[index + 1]
        segment = description[start:end].strip(" :-\n")
        if segment:
            chunks.extend(_split_job_description_segment(segment))
    return chunks


def _split_job_description_segment(segment: str) -> list[str]:
    segments = INLINE_TRAILING_BOILERPLATE_BOUNDARY_RE.split(segment)
    pieces = []
    for subsegment in segments:
        pieces.extend(
            piece.strip(" :-\n")
            for piece in re.split(r"(?<=[.!?])\s+(?=[A-Z0-9])", subsegment)
        )
    chunks: list[str] = []
    for piece in pieces:
        if not piece:
            continue
        chunks.extend(_split_oversized_job_description_piece(piece))
    return chunks


def _split_oversized_job_description_piece(piece: str) -> list[str]:
    if len(piece) <= 900:
        return [piece]
    parts = [part.strip(" :-\n") for part in re.split(r";\s+", piece) if part.strip()]
    if len(parts) > 1:
        return parts
    return [piece]


def _keep_job_description_chunk(
    chunk: str,
    *,
    preserve_company_context: bool = False,
) -> bool:
    text = " ".join(chunk.split())
    if len(text) < JOD_CHUNK_MIN_MEANINGFUL_LENGTH:
        return False

    hard_drop = _contains_any_casefolded(text, JOD_HARD_DROP_MARKERS) or (
        _first_heading_match(text, TRAILING_BOILERPLATE_HEADINGS, strict_single_word=True)
        is not None
    )
    hard_keep = (
        _contains_role_relevant_heading(text)
        or _contains_any_casefolded(text, JOD_HARD_KEEP_MARKERS)
    )
    role_marker_count = _role_relevant_marker_count(text)
    if hard_drop and JOD_BOILERPLATE_START_RE.search(text):
        return False
    if preserve_company_context and not hard_drop:
        return True
    if hard_drop and not hard_keep and role_marker_count < 2:
        return False
    if hard_keep or role_marker_count >= 2:
        return True
    return _JOD_CHUNK_CLASSIFIER.keep_log_odds(text) >= JOD_CHUNK_KEEP_THRESHOLD


def _contains_role_relevant_heading(text: str) -> bool:
    return _first_heading_match(text, ROLE_RELEVANT_START_HEADINGS) is not None


def _role_relevant_marker_count(text: str) -> int:
    normalized = text.casefold()
    return sum(1 for marker in ROLE_RELEVANT_PREFIX_MARKERS if marker in normalized)


def _contains_any_casefolded(text: str, markers: tuple[str, ...]) -> bool:
    normalized = text.casefold()
    return any(marker.casefold() in normalized for marker in markers)


def _jod_chunk_feature_tokens(text: str) -> list[str]:
    return [
        token
        for token in re.findall(r"[a-z][a-z0-9.+#/-]+", text.casefold())
        if len(token) > 2
    ]


_JOD_CHUNK_CLASSIFIER = _JodChunkClassifier.train(JOD_CHUNK_TRAINING_EXAMPLES)


def _first_heading_match(
    text: str,
    headings: tuple[str, ...],
    *,
    start: int = 0,
    skip_before: int = -1,
    strict_single_word: bool = False,
) -> re.Match[str] | None:
    matches = _heading_matches(
        text,
        headings,
        start=start,
        skip_before=skip_before,
        strict_single_word=strict_single_word,
    )
    if not matches:
        return None
    return min(matches, key=lambda match: match.start())


def _heading_matches(
    text: str,
    headings: tuple[str, ...],
    *,
    start: int = 0,
    skip_before: int = -1,
    strict_single_word: bool = False,
) -> list[re.Match[str]]:
    matches: list[re.Match[str]] = []
    for heading in headings:
        for match in _heading_pattern(heading).finditer(text, pos=start):
            if match.start() < skip_before:
                continue
            if _is_heading_like_match(
                text,
                match,
                heading=heading,
                strict_single_word=strict_single_word,
            ):
                matches.append(match)
    return sorted(matches, key=lambda match: match.start())


def _heading_pattern(heading: str) -> re.Pattern[str]:
    parts = [re.escape(part) for part in heading.split()]
    pattern = r"\s+".join(parts)
    return re.compile(rf"(?<![\w/]){pattern}(?=\s|[:?!.,;()\-/–—]|$)", re.IGNORECASE)


def _is_heading_like_match(
    text: str,
    match: re.Match[str],
    *,
    heading: str,
    strict_single_word: bool,
) -> bool:
    matched_text = match.group(0)
    first_alpha = next((char for char in matched_text if char.isalpha()), "")
    if first_alpha and not first_alpha.isupper():
        return False

    heading_word_count = len(heading.split())
    previous_index = match.start() - 1
    while previous_index >= 0 and text[previous_index].isspace():
        previous_index -= 1
    if (
        previous_index >= 0
        and text[previous_index] not in ".!?:;\n\r"
        and heading_word_count == 1
    ):
        return False

    if strict_single_word and heading_word_count == 1:
        suffix = text[match.end() :].lstrip()
        if suffix and suffix[0] not in ":-–—\n\r":
            return False
    return True


def _normalize_job_description_text(description: str) -> str:
    lines = [" ".join(line.split()) for line in description.splitlines()]
    normalized_lines = [line for line in lines if line]
    if len(normalized_lines) > 1:
        return "\n".join(normalized_lines)
    return " ".join(description.split())


def _coerce_search_query(
    value: dict[str, Any],
    *,
    location: str,
    date_posted: DatePosted,
    limit_per_query: int,
) -> JobSearchQuery:
    max_limit = min(max(limit_per_query, 1), 100)
    data = {
        "keywords": str(value.get("keywords") or "").strip(),
        "location": str(value.get("location") or location).strip() or location,
        "date_posted": value.get("date_posted") or date_posted,
        "job_type": value.get("job_type"),
        "workplace_type": value.get("workplace_type"),
        "experience_level": _coerce_search_experience_level(value.get("experience_level")),
        "sort_by": value.get("sort_by") or "recent",
        "distance": value.get("distance"),
        "limit": min(max(int(value.get("limit") or max_limit), 1), max_limit),
        "page": max(int(value.get("page") or 0), 0),
    }
    if not data["keywords"]:
        raise WorkflowError("The LLM returned a search query without keywords.")
    try:
        return JobSearchQuery(**data)
    except ValueError:
        data["date_posted"] = date_posted
        data["job_type"] = None
        data["workplace_type"] = None
        data["experience_level"] = None
        data["sort_by"] = "recent"
        data["distance"] = None
        return JobSearchQuery(**data)


def _coerce_search_experience_level(value: Any) -> str | None:
    normalized = _normalize_experience_level(value)
    if not normalized or normalized in DISALLOWED_EXPERIENCE_LEVELS:
        return None
    return normalized


def _supplement_search_queries(
    queries: list[JobSearchQuery],
    *,
    location: str,
    date_posted: DatePosted,
    limit_per_query: int,
) -> list[JobSearchQuery]:
    max_limit = min(max(limit_per_query, 1), 100)
    supplements = [
        JobSearchQuery(
            keywords=keywords,
            location=location,
            date_posted=date_posted,
            job_type="full_time",
            workplace_type="remote",
            experience_level="mid_senior",
            sort_by="recent",
            limit=max_limit,
            page=0,
        )
        for keywords in DEFAULT_SUPPLEMENTAL_SEARCH_KEYWORDS
    ]

    ordered: list[JobSearchQuery] = []
    seen_keywords: set[str] = set()
    for query in [*queries[:1], *supplements, *queries[1:]]:
        key = query.keywords.casefold().strip()
        if not key or key in seen_keywords:
            continue
        seen_keywords.add(key)
        ordered.append(query)
    return ordered


def _expand_remote_and_hybrid_queries(
    queries: list[JobSearchQuery],
    *,
    max_queries: int,
) -> list[JobSearchQuery]:
    expanded: list[JobSearchQuery] = []
    seen: set[tuple[str, str, str]] = set()
    for query in queries:
        for workplace_type in ("remote", "hybrid"):
            next_query = query.model_copy(update={"workplace_type": workplace_type})
            key = (
                next_query.keywords.casefold(),
                next_query.location.casefold(),
                workplace_type,
            )
            if key in seen:
                continue
            seen.add(key)
            expanded.append(next_query)
            if len(expanded) >= max_queries:
                return expanded
    return expanded


def _search_query_prompt(
    *,
    profile_context: str,
    location: str,
    date_posted: DatePosted,
    limit_per_query: int,
    max_queries: int,
    search_memory: _SearchMemory | None = None,
) -> str:
    feedback_section = ""
    if search_memory and search_memory.total_searches > 0:
        feedback_section = f"""
Search feedback from previous iterations (up to 10 per category):

Keyword strings that returned jobs (reward these — they work):
{search_memory.reward_sample}

Keyword strings that returned zero jobs (avoid these patterns):
{search_memory.penalty_sample}

Use this feedback to generate fresh, non-duplicate keyword combinations.
Prefer short, general keyword strings (2-4 words) over long, specific ones.
Include role titles like "Platform Engineer", "DevOps Engineer", "Software Engineer",
combined with domain terms like "infrastructure", "automation", "cloud", "distributed systems",
"AI", "agentic AI", and "LLMs" when supported by the profile.
"""

    return f"""
You generate LinkedIn public job search parameters from a candidate profile.
Return only valid JSON. Do not include commentary.
{feedback_section}
Rules:
- Generate up to {max_queries} keyword-focused query objects.
- Use combinations of role titles, seniority, domain keywords, and core skills.
- The workflow will force remote and hybrid LinkedIn workplace filters.
- Do not use internship or entry_level experience filters.
- Prefer concise keyword strings that LinkedIn search can use directly.
- Include AI, agentic AI, and LLM keyword variants when they overlap with the candidate profile.
- Do not invent facts about the candidate.

Each query object must use this schema:
{{
  "keywords": "string",
  "location": "{location}",
  "date_posted": "{date_posted}",
  "job_type": "full_time",
  "workplace_type": "remote",
  "experience_level": "mid_senior",
  "sort_by": "recent",
  "limit": {limit_per_query},
  "page": 0
}}

Allowed values:
- date_posted: any_time, past_24_hours, past_week, past_month
- job_type: full_time, part_time, contract, temporary, volunteer, internship, other
- workplace_type: remote, hybrid
- experience_level: {", ".join(SEARCH_EXPERIENCE_LEVELS)}
- sort_by: relevance, recent

Candidate profile files:
{profile_context}
""".strip()


def _scjdir_prompt(
    *,
    source_resume: ProfileDocument | None,
    current_job_description: ProfileDocument | None,
    job: JobDetails,
) -> str:
    source_resume_text = _limit_context(
        source_resume.text if source_resume else "",
        max_chars=12_000,
    )
    cjd_text = _limit_context(
        current_job_description.text if current_job_description else "No CJD was available.",
        max_chars=8_000,
    )
    return f"""
You are rewriting only the candidate's current-role resume section, called SCJDiR.
Return only the replacement SCJDiR block. Do not return advice, notes, markdown fences,
JSON, analysis, or instructions.

Rules:
- Keep the employer, location, title, and dates factual.
- Use the CJD only as supporting context for the current Oracle role.
- Use the job opening description only to choose emphasis and language.
- Make small, factual wording changes around the margins.
- Preserve the approximate length and bullet count of the original SCJDiR.
- Use plain resume lines only; do not prefix lines with bullets, hyphens, or bullet glyphs.
- Do not invent products, dates, employers, certifications, tools, metrics, or responsibilities.
- Prefer resume bullets, not recommendations.

Original SCJDiR:
{DEFAULT_SCJDIR}

Current resume text:
{source_resume_text}

Current job description (CJD):
{cjd_text}

Job opening description (JOD):
Title: {job.title}
Company: {job.company or "Unknown"}
Location: {job.location or "Unknown"}
LinkedIn job ID: {job.job_id}
Description:
{_job_description_context(job)}
""".strip()


def _resume_sections_prompt(
    *,
    source_resume: ProfileDocument | None,
    current_job_description: ProfileDocument | None,
    tailored_scjdir: str,
    job: JobDetails,
) -> str:
    source_resume_text = _limit_context(
        source_resume.text if source_resume else "",
        max_chars=18_000,
    )
    cjd_hint = _limit_context(
        current_job_description.text if current_job_description else "No CJD was available.",
        max_chars=6_000,
    )
    default_skills = "\n".join(
        f"- {category}: {', '.join(skills)}"
        for category, skills in DEFAULT_CORE_TECHNICAL_SKILLS
    )
    default_prior_experience = _render_prior_experience_text(DEFAULT_PRIOR_EXPERIENCE_ENTRIES)
    return f"""
You produce structured resume section edits for one job opening.
Return only valid JSON. Do not return markdown fences, commentary, advice, or a full resume.

The application will render the final resume from a local template. You only control:
1. core_technical_skills
2. prior_experience

Do not include the header, professional summary, AI generation note, Oracle current role, or
education/certifications in the JSON response. Those sections are static or generated separately.

Hard requirements:
- Preserve the seven Core Technical Skills categories exactly.
- Add or remove individual skills only when supported by the source resume, CJD, or tailored SCJDiR.
- Prefer skills that overlap with the JOD, including AI, agentic AI, and LLM terms only
  when factual.
- Always keep the AI Tools category and use it for AI-assisted engineering tools such as
  Codex, Oracle Code Assist (OCA), Cline, OpenRouter, ChatGPT, or LLM prompting.
- Do not include "Error Budgets" in Core Technical Skills; use concrete observability tools
  or dashboard/pipeline skills instead.
- Preserve all prior employers, locations, titles, dates, and the original bullet count per job.
- For prior experience, make only minor keyword swaps or wording changes that remain factual.
- Do not invent employers, dates, credentials, projects, tools, metrics, or responsibilities.

Return this exact JSON shape:
{{
  "core_technical_skills": [
    {{"category": "Languages & Frameworks", "skills": ["Python", "Django"]}}
  ],
  "prior_experience": [
    {{
      "organization": "University of Iowa Hospitals and Clinics",
      "bullets": ["Full-Cycle Software Engineering: ..."]
    }}
  ]
}}

Tailored Oracle current-role SCJDiR, already generated separately:
{tailored_scjdir}

Base Core Technical Skills:
{default_skills}

Base prior experience section:
{default_prior_experience}

Source resume:
{source_resume_text}

CJD context:
{cjd_hint}

JOD:
Title: {job.title}
Company: {job.company or "Unknown"}
Description:
{_job_description_context(job)}
""".strip()


def _ats_resume_repair_prompt(
    *,
    source_evidence: Mapping[str, str],
    current_resume_text: str,
    current_tailored_scjdir: str,
    current_sections_plan: Mapping[str, Any],
    job: JobDetails,
    score: AtsProxyScore,
) -> str:
    missing_terms = ", ".join(score.missing_high_value_terms[:ATS_REPAIR_MAX_MISSING_TERMS])
    evidence_text = "\n".join(
        f"- {term}: {snippet}" for term, snippet in source_evidence.items()
    )
    current_sections_json = json.dumps(
        current_sections_plan,
        ensure_ascii=True,
        default=str,
        indent=2,
    )
    return f"""
You repair a generated resume after local ATS scoring.
Return only valid JSON. Do not return markdown fences, commentary, advice, or a full resume.

The first resume draft already used the source resume, CJD, and JOD. To save tokens, this
repair pass only uses the JOD, current draft, ATS score, and source-resume evidence snippets.

Goal:
- Improve the ATS score toward {ATS_REPAIR_TARGET_SCORE}/100 when factual source evidence
  supports it.
- Cover missing high-value JOD terms only when supported by the source-resume evidence below.
- Keep the existing local resume template; you only control tailored_scjdir, core_technical_skills,
  and prior_experience.

Rules:
- Do not invent employers, dates, credentials, projects, tools, metrics, products, or
  responsibilities.
- Do not add a missing term unless it appears in, or is directly supported by, the source evidence.
- Prefer small factual keyword/wording changes over broad rewrites.
- Preserve the seven Core Technical Skills categories exactly.
- Preserve prior employer names, locations, titles, dates, and bullet count per job.
- Do not include "Error Budgets" in Core Technical Skills; use concrete observability tools
  or reliability language instead.
- If no factual repair is possible, return the current values unchanged.

Return this JSON shape. Include complete replacement values only for sections you change:
{{
  "tailored_scjdir": "Oracle | Remote / International Datacenters...",
  "core_technical_skills": [
    {{"category": "Languages & Frameworks", "skills": ["Python", "Django"]}}
  ],
  "prior_experience": [
    {{
      "organization": "University of Iowa Hospitals and Clinics",
      "bullets": ["Full-Cycle Software Engineering: ..."]
    }}
  ]
}}

ATS score:
- overall: {score.overall_score}/100
- parsing: {score.parsing_score}/100
- keyword: {score.keyword_match_score}/100
- semantic: {score.semantic_match_score}/100
- formatting risk: {score.formatting_risk}
- missing high-value terms: {missing_terms or "None"}

Source-resume evidence for missing terms:
{evidence_text}

Current tailored Oracle section:
{_limit_context(current_tailored_scjdir, max_chars=4_000)}

Current dynamic section plan:
{_limit_context(current_sections_json, max_chars=5_000)}

Current generated resume:
{_limit_context(current_resume_text, max_chars=10_000)}

JOD:
Title: {job.title}
Company: {job.company or "Unknown"}
Description:
{_limit_context(_job_description_context(job), max_chars=6_000)}
""".strip()


def _recommendations_prompt(
    *,
    source_resume: ProfileDocument | None,
    current_job_description: ProfileDocument | None,
    job: JobDetails,
    draft_text: str,
) -> str:
    source_resume_text = _limit_context(
        source_resume.text if source_resume else "",
        max_chars=14_000,
    )
    cjd_text = _limit_context(
        current_job_description.text if current_job_description else "No CJD was available.",
        max_chars=6_000,
    )
    return f"""
The prior generation did not produce a usable resume. Produce focused implementation
recommendations that another model or Codex skill can apply to the source resume.

Return only concise recommendations. Do not write a full resume.

Required format:
1. SCJDiR replacement: provide the exact replacement Oracle current-role section.
2. Other minor resume edits: list only small keyword or wording changes outside SCJDiR.
3. Do-not-change constraints: list factual details that must remain unchanged.

Source resume:
{source_resume_text}

CJD:
{cjd_text}

JOD:
Title: {job.title}
Company: {job.company or "Unknown"}
Description:
{_job_description_context(job)}

Unusable draft text:
{_limit_context(draft_text, max_chars=6_000)}
""".strip()


def _cover_letter_sections_prompt(
    *,
    source_resume: ProfileDocument | None,
    current_job_description: ProfileDocument | None,
    job: JobDetails,
) -> str:
    source_resume_text = _limit_context(
        source_resume.text if source_resume else "",
        max_chars=14_000,
    )
    cjd_text = _limit_context(
        current_job_description.text if current_job_description else "No CJD was available.",
        max_chars=8_000,
    )
    prior_experience_text = _render_prior_experience_text(DEFAULT_PRIOR_EXPERIENCE_ENTRIES)
    return f"""
You produce dynamic cover_letter_sections JSON for a local cover-letter template.
Return only valid JSON. Do not return markdown fences, commentary, a full cover letter, or
instructions.

The application renders these static sections itself:
- Date, salutation, resume-attached sentence, thank-you sentence, and signature.
- Section 1 opening sentence with the real job title and company.
- Section 2 starts with this exact Oracle sentence:
  {COVER_LETTER_ORACLE_OPENER}
- Section 3 starts with this exact sentence:
  {COVER_LETTER_PRIOR_EXPERIENCE_OPENER}
- Section 4 is a static paragraph about the candidate's linkedin-career-mcp automation project,
  OpenRouter/DeepSeek, Codex, Cline with DeepSeek, GitHub Copilot, prompt structure, context
  windows, token usage, model selection, and output validation.

You only control:
1. opening_alignment: a concise noun phrase for the end of Section 1.
2. oracle_alignment: exactly 3 concise sentences connecting the Oracle current-role resume
   section from the source resume, the current Oracle job description/CJD, and this JOD.
3. prior_experience_alignment: 2-3 concise sentences connecting pre-Oracle experience to this JOD.

Rules:
- Stay factual. Do not invent employers, projects, credentials, tools, dates, metrics, products,
  leadership scope, or submitted-application facts.
- For oracle_alignment, use the source resume's Oracle current-role section for proven
  accomplishments, use the CJD for richer current-role context, and use the JOD to choose which
  accomplishments to emphasize.
- Each oracle_alignment sentence should add a distinct alignment point around platform automation,
  IaC/CI/CD, observability, API integration, developer tooling, distributed systems, cloud
  infrastructure, or AI-assisted engineering when relevant.
- Use the JOD to choose emphasis and vocabulary.
- Use the CJD only to describe the current Oracle role.
- Use the prior experience context for pre-Oracle jobs; do not turn Oracle facts into prior jobs.
- Avoid flattery, buzzword stuffing, and generic claims.
- Do not mention salary, availability, relocation, immigration, or application logistics.
- Do not include greetings, closings, markdown, bullets, or section labels.

Return this exact JSON shape:
{{
  "opening_alignment": "the platform and AI automation capabilities you are looking for",
  "oracle_alignment": "At Oracle, I have ... . That work also ... . This maps to the role ... .",
  "prior_experience_alignment": "At the University of Iowa Hospitals and Clinics, ..."
}}

Source resume:
{source_resume_text}

Current job description (CJD):
{cjd_text}

Pre-Oracle experience context from the resume:
{prior_experience_text}

Job opening description (JOD):
Title: {job.title}
Company: {job.company or "Unknown"}
Location: {job.location or "Unknown"}
LinkedIn job ID: {job.job_id}
Description:
{_job_description_context(job)}
""".strip()


def _render_cover_letter_template(
    *,
    job: JobDetails,
    sections_plan: Mapping[str, Any],
    letter_date: date | None = None,
) -> str:
    sections = _coerce_cover_letter_sections(sections_plan)
    company = job.company or "[Company Name]"
    title = job.title or "[Position Title]"
    opening = (
        f"Please accept my application for the position of {title} at {company}. Having read "
        "through the job description, I am excited to apply because my previous experience and "
        f"technical background align closely with {sections['opening_alignment']}."
    )
    oracle = f"{COVER_LETTER_ORACLE_OPENER} {sections['oracle_alignment']}"
    prior = (
        f"{COVER_LETTER_PRIOR_EXPERIENCE_OPENER} "
        f"{sections['prior_experience_alignment']}"
    )

    lines = [
        _format_cover_letter_date(letter_date or date.today()),
        "",
        "Dear Hiring Manager,",
        "",
        opening,
        "",
        oracle,
        "",
        prior,
        "",
        COVER_LETTER_PROJECT_PARAGRAPH,
        "",
        "Please find my resume attached for your consideration.",
        "",
        "Thank you for your time. I look forward to hearing from you and would be happy to "
        "answer any questions.",
        "",
        "Sincerely,",
        "Maxim Perkhounkov",
        LINKEDIN_PROFILE_MARKDOWN,
    ]
    return _clean_cover_letter_text("\n".join(lines))


def _coerce_cover_letter_sections(raw_value: Any) -> dict[str, str]:
    if not isinstance(raw_value, Mapping):
        raw_value = {}
    return {
        "opening_alignment": _clean_cover_letter_fragment(
            raw_value.get("opening_alignment"),
            default=COVER_LETTER_DEFAULT_OPENING_ALIGNMENT,
            max_chars=220,
            ensure_sentence=False,
        ),
        "oracle_alignment": _clean_cover_letter_fragment(
            raw_value.get("oracle_alignment"),
            default=COVER_LETTER_DEFAULT_ORACLE_ALIGNMENT,
            max_chars=1050,
            ensure_sentence=True,
        ),
        "prior_experience_alignment": _clean_cover_letter_fragment(
            raw_value.get("prior_experience_alignment"),
            default=COVER_LETTER_DEFAULT_PRIOR_EXPERIENCE_ALIGNMENT,
            max_chars=950,
            ensure_sentence=True,
        ),
    }


def _clean_cover_letter_fragment(
    value: Any,
    *,
    default: str,
    max_chars: int,
    ensure_sentence: bool,
) -> str:
    text = (
        " ".join(str(item) for item in value)
        if isinstance(value, list)
        else str(value or "")
    )
    text = _strip_markdown_emphasis(_clean_resume_text(text))
    text = re.sub(
        r"^\s*(?:opening_alignment|oracle_alignment|prior_experience_alignment)\s*:\s*",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = _clean_inline_text(text)
    if not text:
        text = default
    text = text[:max_chars].strip(" ,;")
    if not ensure_sentence:
        return text.rstrip(".!?")
    if text and text[-1] not in ".!?":
        text = f"{text}."
    return text


def _format_cover_letter_date(value: date) -> str:
    return f"{value.strftime('%B')} {value.day}, {value.year}"


def _clean_cover_letter_text(text: str) -> str:
    lines = [_clean_inline_text(line) for line in text.splitlines()]
    cleaned: list[str] = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank and cleaned:
                cleaned.append("")
            previous_blank = True
            continue
        cleaned.append(line)
        previous_blank = False
    return "\n".join(cleaned).strip()


def _coerce_ats_repaired_resume_draft(
    *,
    repair_plan: Any,
    current_draft: _ResumeDraft,
) -> _ResumeDraft:
    if not isinstance(repair_plan, Mapping):
        return current_draft

    tailored_scjdir = current_draft.tailored_scjdir
    raw_scjdir = repair_plan.get("tailored_scjdir") or repair_plan.get("scjdir")
    if isinstance(raw_scjdir, str) and _resume_block_lines(raw_scjdir):
        tailored_scjdir = raw_scjdir

    sections_plan: dict[str, Any] = dict(current_draft.sections_plan)
    core_skills = _merge_category_section_plan(
        current_value=sections_plan.get("core_technical_skills"),
        repair_value=repair_plan.get("core_technical_skills"),
        category_key="category",
    )
    if core_skills is not None:
        sections_plan["core_technical_skills"] = core_skills

    prior_experience = _merge_category_section_plan(
        current_value=sections_plan.get("prior_experience"),
        repair_value=repair_plan.get("prior_experience"),
        category_key="organization",
    )
    if prior_experience is not None:
        sections_plan["prior_experience"] = prior_experience

    text = _render_resume_template(
        tailored_scjdir=tailored_scjdir,
        sections_plan=sections_plan,
    )
    return _ResumeDraft(
        tailored_scjdir=tailored_scjdir,
        sections_plan=sections_plan,
        text=text,
    )


def _merge_category_section_plan(
    *,
    current_value: Any,
    repair_value: Any,
    category_key: str,
) -> list[Any] | None:
    if not isinstance(repair_value, list):
        return None
    if not isinstance(current_value, list):
        return repair_value

    merged_by_key: dict[str, Any] = {}
    order: list[str] = []
    for item in [*current_value, *repair_value]:
        if not isinstance(item, Mapping):
            continue
        raw_key = str(item.get(category_key) or item.get("name") or item.get("company") or "")
        normalized_key = _normalize_label(raw_key)
        if not normalized_key:
            continue
        if normalized_key not in order:
            order.append(normalized_key)
        merged_by_key[normalized_key] = dict(item)
    return [merged_by_key[key] for key in order]


def _render_resume_template(
    *,
    tailored_scjdir: str,
    sections_plan: Mapping[str, Any],
) -> str:
    core_skills = _coerce_core_skill_sections(sections_plan.get("core_technical_skills"))
    prior_experience = _coerce_prior_experience_entries(sections_plan.get("prior_experience"))
    scjdir_lines = _resume_block_lines(tailored_scjdir)
    if not scjdir_lines or _looks_like_recommendations("\n".join(scjdir_lines)):
        scjdir_lines = _resume_block_lines(DEFAULT_SCJDIR)

    lines: list[str] = [
        RESUME_HEADER_NAME,
        RESUME_HEADER_CONTACT,
        "",
        "Professional Summary",
        STATIC_PROFESSIONAL_SUMMARY,
        "",
        AI_GENERATION_NOTE,
        "",
        "Core Technical Skills",
    ]
    lines.extend(_render_core_skills_lines(core_skills))
    lines.extend(["", "Professional Experience"])
    lines.extend(scjdir_lines)
    lines.append("")
    lines.extend(_render_prior_experience_lines(prior_experience))
    lines.extend(["", "Education & Certifications"])
    lines.extend(DEFAULT_EDUCATION_CERTIFICATIONS)
    return "\n".join(lines).strip()


def _coerce_core_skill_sections(raw_value: Any) -> list[tuple[str, list[str]]]:
    default_sections = [
        (category, list(skills)) for category, skills in DEFAULT_CORE_TECHNICAL_SKILLS
    ]
    if not isinstance(raw_value, list):
        return default_sections

    by_category: dict[str, list[str]] = {}
    for item in raw_value:
        category = ""
        skills_value: Any = None
        if isinstance(item, Mapping):
            category = str(item.get("category") or item.get("name") or "").strip()
            skills_value = item.get("skills") or item.get("items")
        elif isinstance(item, str) and ":" in item:
            category, skills_value = item.split(":", 1)
        skills = _coerce_skill_items(skills_value)
        if category and skills:
            by_category[_normalize_label(category)] = skills

    sections: list[tuple[str, list[str]]] = []
    for category, default_skills in default_sections:
        skills = by_category.get(_normalize_label(category), default_skills)
        if category == "AI Tools":
            skills = [*skills, *default_skills]
        clean_skills = _dedupe_preserve_order(
            [skill for skill in skills if _is_allowed_core_skill(skill)]
        )
        sections.append((category, clean_skills[:12] or list(default_skills)))
    return sections


def _coerce_skill_items(value: Any) -> list[str]:
    if isinstance(value, str):
        raw_items = re.split(r",|;|\n", value)
    elif isinstance(value, list):
        raw_items = [str(item) for item in value]
    else:
        return []
    items: list[str] = []
    for item in raw_items:
        clean_item = _clean_list_item_text(item)
        if clean_item and _is_allowed_core_skill(clean_item):
            items.append(clean_item)
    return items


def _coerce_prior_experience_entries(raw_value: Any) -> list[dict[str, object]]:
    default_entries = [
        {
            "organization": entry["organization"],
            "location": entry["location"],
            "title": entry["title"],
            "dates": entry["dates"],
            "bullets": list(entry["bullets"]),
        }
        for entry in DEFAULT_PRIOR_EXPERIENCE_ENTRIES
    ]
    if not isinstance(raw_value, list):
        return default_entries

    raw_by_org: dict[str, Mapping[str, Any]] = {}
    for item in raw_value:
        if not isinstance(item, Mapping):
            continue
        organization = str(item.get("organization") or item.get("company") or "").strip()
        if organization:
            raw_by_org[_normalize_label(organization)] = item

    entries: list[dict[str, object]] = []
    for default_entry in default_entries:
        raw_entry = raw_by_org.get(_normalize_label(str(default_entry["organization"])))
        if raw_entry is None:
            entries.append(default_entry)
            continue

        default_bullets = list(default_entry["bullets"])
        model_bullets = _coerce_bullet_items(raw_entry.get("bullets"))
        bullet_count = len(default_bullets)
        if not model_bullets:
            bullets = default_bullets
        else:
            bullets = [
                model_bullets[index] if index < len(model_bullets) else default_bullets[index]
                for index in range(bullet_count)
            ]
        entries.append({**default_entry, "bullets": bullets})
    return entries


def _coerce_bullet_items(value: Any) -> list[str]:
    if isinstance(value, str):
        raw_items = [line for line in value.splitlines() if line.strip()]
    elif isinstance(value, list):
        raw_items = [str(item) for item in value]
    else:
        return []
    return [
        clean_item
        for item in raw_items
        if (clean_item := _clean_list_item_text(item))
    ]


def _render_core_skills_lines(sections: list[tuple[str, list[str]]]) -> list[str]:
    return [f"- {category}: {', '.join(skills)}" for category, skills in sections]


def _render_prior_experience_lines(entries: list[dict[str, object]]) -> list[str]:
    lines: list[str] = []
    for index, entry in enumerate(entries):
        if index:
            lines.append("")
        lines.append(f"{entry['organization']} | {entry['location']}")
        lines.append(f"{entry['title']} | {entry['dates']}")
        lines.extend(f"- {bullet}" for bullet in entry["bullets"])
    return lines


def _render_prior_experience_text(entries: tuple[dict[str, object], ...]) -> str:
    return "\n".join(_render_prior_experience_lines(list(entries)))


def _resume_block_lines(text: str) -> list[str]:
    lines = []
    for raw_line in _clean_resume_text(text).splitlines():
        line = _clean_inline_text(raw_line)
        if not line or line in RESUME_SECTION_HEADINGS:
            continue
        line_body, had_bullet_marker = _strip_leading_bullet_markers(line)
        if not line_body or line_body in RESUME_SECTION_HEADINGS:
            continue
        if _looks_like_model_intro_line(line_body):
            continue
        if len(lines) >= 2 and (had_bullet_marker or ":" in line_body):
            line_body = f"- {line_body}"
        lines.append(line_body)
    return lines


def _clean_inline_text(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _clean_list_item_text(text: str) -> str:
    return _strip_leading_bullet_markers(_clean_inline_text(text))[0]


def _looks_like_model_intro_line(line: str) -> bool:
    normalized = line.casefold().strip()
    intro_prefixes = (
        "here is ",
        "here's ",
        "here are ",
        "below is ",
        "below are ",
        "the rewritten ",
        "rewritten ",
    )
    resume_markers = ("scjdir", "resume", "section", "version")
    return normalized.startswith(intro_prefixes) and any(
        marker in normalized for marker in resume_markers
    )


def _strip_leading_bullet_markers(text: str) -> tuple[str, bool]:
    line = text.strip()
    had_bullet_marker = False
    while True:
        match = LEADING_BULLET_RE.match(line)
        if not match:
            return line, had_bullet_marker
        had_bullet_marker = True
        line = line[match.end() :].strip()


def _is_allowed_core_skill(skill: str) -> bool:
    normalized = re.sub(r"[^a-z0-9]+", " ", skill.casefold()).strip()
    return normalized not in DISALLOWED_CORE_SKILLS and "error budget" not in normalized


def _normalize_label(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.casefold())


def _dedupe_preserve_order(values: list[str]) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for value in values:
        key = value.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(value)
    return result


def _glob_matches(value: str, pattern: str) -> bool:
    regex = "^" + re.escape(pattern).replace(r"\*", ".*").replace(r"\?", ".") + "$"
    return re.match(regex, value) is not None


def _merge_artifacts(
    existing: TailoredResumeArtifact | None,
    update: TailoredResumeArtifact,
) -> TailoredResumeArtifact:
    if existing is None:
        return update
    return TailoredResumeArtifact(
        job_id=existing.job_id,
        company=existing.company or update.company,
        title=existing.title or update.title,
        linkedin_url=existing.linkedin_url or update.linkedin_url,
        resume_path=existing.resume_path or update.resume_path,
        cover_letter_path=existing.cover_letter_path or update.cover_letter_path,
        artifact_kind=existing.artifact_kind
        if existing.artifact_kind != "cover_letter"
        else update.artifact_kind,
        recommendations_path=existing.recommendations_path or update.recommendations_path,
    )


def audit_application_artifacts(database_path: Path) -> ArtifactAudit:
    with connect_database(database_path) as connection:
        rows = connection.execute(
            """
            SELECT job_id, company, job_title,
                   resume_content IS NOT NULL AS has_resume,
                   cover_letter_content IS NOT NULL AS has_cover_letter
            FROM applications
            ORDER BY company COLLATE NOCASE ASC, job_title COLLATE NOCASE ASC
            """
        ).fetchall()

    missing_artifacts: list[MissingApplicationArtifact] = []
    with_resumes = 0
    with_cover_letters = 0
    for row in rows:
        missing: list[Literal["resume", "cover_letter"]] = []
        if row["has_resume"]:
            with_resumes += 1
        else:
            missing.append("resume")
        if row["has_cover_letter"]:
            with_cover_letters += 1
        else:
            missing.append("cover_letter")
        if missing:
            missing_artifacts.append(
                MissingApplicationArtifact(
                    job_id=str(row["job_id"]),
                    company=str(row["company"] or ""),
                    title=str(row["job_title"] or ""),
                    missing=missing,
                )
            )

    return ArtifactAudit(
        total_jobs=len(rows),
        with_resumes=with_resumes,
        with_cover_letters=with_cover_letters,
        missing_resumes=sum(
            1 for artifact in missing_artifacts if "resume" in artifact.missing
        ),
        missing_cover_letters=sum(
            1 for artifact in missing_artifacts if "cover_letter" in artifact.missing
        ),
        missing_artifacts=missing_artifacts,
    )


def _emit_job_progress(
    progress_callback: ProgressCallback | None,
    *,
    action: str,
    job: JobDetails,
    index: int,
    total: int,
    detail: str | None = None,
) -> None:
    if progress_callback is None:
        return
    prefix = f"{action} {index}/{total}"
    if detail:
        prefix = f"{prefix} ({detail})"
    progress_callback(
        f"{prefix} - job title: '{job.title}', "
        f"company: '{job.company or 'Unknown'}', job id: '{job.job_id}'"
    )


def _emit_progress(progress_callback: ProgressCallback | None, message: str) -> None:
    if progress_callback is not None:
        progress_callback(message)


def _stderr_progress(message: str) -> None:
    print(message, file=sys.stderr, flush=True)


class _WorkflowRunLogger:
    def __init__(self, *, output_dir: Path, operation: str, llm_label: str) -> None:
        timestamp = datetime.now(UTC).strftime("%Y%m%dT%H%M%SZ")
        self.path = output_dir / "logs" / f"{timestamp}_{operation}.jsonl"
        self.path.parent.mkdir(parents=True, exist_ok=True)
        self.write("start", operation=operation, llm=llm_label)

    def progress(self, message: str) -> None:
        _stderr_progress(message)
        self.write("progress", message=message)

    def result(self, result: MatchingJobsWorkflowResult) -> None:
        self.write(
            "result",
            jobs_found=result.jobs_found,
            resumes_created=result.resumes_created,
            cover_letters_created=result.cover_letters_created,
            recommendations_created=result.recommendations_created,
            errors=result.errors,
            artifact_audit=result.artifact_audit.model_dump(mode="json"),
        )

    def failure(self, exc: BaseException) -> None:
        self.write("failure", error_type=type(exc).__name__, error=str(exc))

    def write(self, event: str, **fields: object) -> None:
        record = {
            "timestamp": datetime.now(UTC).isoformat(timespec="seconds"),
            "event": event,
            **fields,
        }
        with self.path.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(record, ensure_ascii=True) + "\n")


def _job_label(company: str | None, title: str | None) -> str:
    return f"{company or 'Unknown company'} - {title or 'Unknown title'}"


def _path_part(value: str, *, lower: bool = False) -> str:
    value = value.lower() if lower else value
    value = re.sub(r"[^A-Za-z0-9._ -]+", "", value)
    value = re.sub(r"[\s/]+", "_", value.strip())
    value = value.strip("._-")
    return value[:120] or "unknown"


def _clean_resume_text(text: str) -> str:
    text = re.sub(r"```(?:\w+)?", "", text)
    return text.replace("```", "").strip()


INLINE_MARKUP_RE = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)|\*\*(.+?)\*\*")


def _paragraph_markup(line: str) -> str:
    parts: list[str] = []
    cursor = 0
    for match in INLINE_MARKUP_RE.finditer(line):
        parts.append(escape(line[cursor : match.start()]))
        if match.group(1) is not None:
            label = escape(match.group(1))
            url = escape(match.group(2), quote=True)
            parts.append(f'<a href="{url}" color="blue">{label}</a>')
        else:
            parts.append(f"<b>{_paragraph_markup(match.group(3))}</b>")
        cursor = match.end()
    parts.append(escape(line[cursor:]))
    return "".join(parts)


def _cover_letter_markup(paragraph: str) -> str:
    return "<br/>".join(
        _paragraph_markup(line.strip()) for line in paragraph.splitlines() if line.strip()
    )


def _bullet_markup(line: str) -> str:
    line = _strip_markdown_emphasis(line)
    label_parts = _split_unprotected_separator(line, ":")
    if label_parts is not None:
        label, rest = label_parts
        if 2 <= len(label) <= 60:
            if _has_markdown_link(label):
                return f"{_paragraph_markup(label)}:{_paragraph_markup(rest)}"
            return f"<b>{_paragraph_markup(label)}:</b>{_paragraph_markup(rest)}"
    label_parts = _split_unprotected_separator(line, "|")
    if label_parts is not None:
        label, rest = label_parts
        if 2 <= len(label) <= 80:
            if _has_markdown_link(label):
                return f"{_paragraph_markup(label.strip())} |{_paragraph_markup(rest)}"
            return f"<b>{_paragraph_markup(label.strip())}</b> |{_paragraph_markup(rest)}"
    return _paragraph_markup(line)


def _nested_bullet_markup(line: str) -> str:
    line = _strip_markdown_emphasis(line)
    label_parts = _split_unprotected_separator(line, ":")
    if label_parts is not None:
        label, rest = label_parts
        if 2 <= len(label) <= 60:
            if _has_markdown_link(label):
                return f"{_paragraph_markup(label)}:{_paragraph_markup(rest)}"
            return f"<i>{_paragraph_markup(label)}:</i>{_paragraph_markup(rest)}"
    return _paragraph_markup(line)


def _title_markup(line: str) -> str:
    line = _strip_markdown_emphasis(line)
    title_parts = _split_unprotected_separator(line, "|")
    if title_parts is None:
        return _paragraph_markup(line)
    title, dates = title_parts
    if _has_markdown_link(title):
        return f"{_paragraph_markup(title.strip())} | <i>{_paragraph_markup(dates.strip())}</i>"
    return f"<b>{_paragraph_markup(title.strip())}</b> | <i>{_paragraph_markup(dates.strip())}</i>"


def _split_unprotected_separator(line: str, separator: str) -> tuple[str, str] | None:
    protected_spans = [match.span() for match in re.finditer(r"\[[^\]]+]\(https?://[^)]+\)", line)]
    for index, character in enumerate(line):
        if character != separator:
            continue
        if any(start <= index < end for start, end in protected_spans):
            continue
        return line[:index], line[index + 1 :]
    return None


def _has_markdown_link(text: str) -> bool:
    return bool(re.search(r"\[[^\]]+]\(https?://[^)]+\)", text))


def _strip_markdown_emphasis(text: str) -> str:
    return re.sub(r"(?<!\*)\*\*(.+?)\*\*(?!\*)", r"\1", text)


def _cover_letter_top_bar() -> HRFlowable:
    return HRFlowable(
        width="100%",
        thickness=7.2,
        color=EMERALD_ACCENT,
        spaceBefore=0,
        spaceAfter=0,
    )


def _resume_rule(*, top: bool = False) -> HRFlowable:
    return HRFlowable(
        width="100%",
        thickness=2.0 if top else 1.35,
        color=EMERALD_ACCENT,
        spaceBefore=0,
        spaceAfter=0,
    )


def _looks_like_employer_line(line: str) -> bool:
    if line.startswith("- ") or " | " not in line:
        return False
    return not any(date in line for date in ("Jan ", "Feb ", "Mar ", "Apr ", "May ", "Jun "))


def _looks_like_title_line(line: str) -> bool:
    if line.startswith("- ") or " | " not in line:
        return False
    return any(
        marker in line
        for marker in (
            "Present",
            "Jan ",
            "Feb ",
            "Mar ",
            "Apr ",
            "May ",
            "Jun ",
            "Jul ",
            "Aug ",
            "Sep ",
            "Oct ",
            "Nov ",
            "Dec ",
        )
    )


def _looks_like_recommendations(text: str) -> bool:
    normalized = text.casefold()
    recommendation_markers = (
        "recommendation",
        "recommendations",
        "suggestion",
        "suggestions",
        "here are",
        "you should",
        "i recommend",
        "changes to make",
        "proposed changes",
    )
    has_marker = any(marker in normalized for marker in recommendation_markers)
    has_resume_anchor = "professional summary" in normalized or "oracle" in normalized
    return has_marker and not has_resume_anchor


def _looks_like_heading(line: str) -> bool:
    if len(line) > 80:
        return False
    letters = [char for char in line if char.isalpha()]
    return bool(letters) and sum(char.isupper() for char in letters) / len(letters) > 0.75


def _looks_like_extra_resume_section_heading(line: str, *, current_section: str | None) -> bool:
    if current_section is None or current_section == "Professional Experience":
        return False
    if (
        current_section in RESUME_SECTION_HEADINGS
        and current_section != "Education & Certifications"
    ):
        return False
    if line in RESUME_SECTION_HEADINGS or line.startswith(("- ", "  - ", "Note:")):
        return False
    if len(line) > 80 or any(separator in line for separator in (":", "|")):
        return False
    if line.endswith((".", ",", ";")):
        return False

    words = re.findall(r"[A-Za-z][A-Za-z0-9+/.-]*", line)
    if len(words) < 2:
        return False
    lowercase_allowed = {
        "a",
        "an",
        "and",
        "as",
        "for",
        "in",
        "of",
        "on",
        "or",
        "the",
        "to",
        "with",
    }
    meaningful_words = [word for word in words if word.casefold() not in lowercase_allowed]
    if len(meaningful_words) < 2:
        return False
    return all(word[:1].isupper() or word.isupper() for word in meaningful_words)


def _ensure_tracking_headers(sheet: Worksheet) -> None:
    existing = [
        sheet.cell(row=1, column=column).value
        for column in range(1, len(TRACKING_HEADERS) + 1)
    ]
    if existing == TRACKING_HEADERS:
        return
    legacy_existing = [
        sheet.cell(row=1, column=column).value
        for column in range(1, len(LEGACY_TRACKING_HEADERS) + 1)
    ]
    if legacy_existing == LEGACY_TRACKING_HEADERS:
        sheet.insert_cols(6, 1)
        sheet.cell(row=1, column=6, value="cover_letter")
        _size_tracking_columns(sheet)
        return
    if sheet.max_row == 1 and all(value is None for value in existing):
        for column, header in enumerate(TRACKING_HEADERS, start=1):
            sheet.cell(row=1, column=column, value=header)
        _size_tracking_columns(sheet)
        return
    raise WorkflowError(f"Tracking workbook has unexpected headers: {sheet.title}")


def _size_tracking_columns(sheet: Worksheet) -> None:
    widths = [16, 24, 36, 60, 70, 70, 14, 16]
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + index)].width = width


def _build_llm_client(
    settings: Settings,
    *,
    api_model: str | None = None,
) -> ApiLlmClient | OllamaClient:
    """Build the LLM client based on the configured provider.

    Defaults to the external API (OpenRouter / DeepSeek). Local Ollama is
    only used when explicitly requested with LINKEDIN_CAREER_MCP_LLM_PROVIDER=ollama.
    """
    provider = settings.llm_provider.casefold().strip()
    if provider == "ollama":
        return OllamaClient(
            base_url=settings.ollama_base_url,
            model=settings.ollama_model,
            timeout_seconds=settings.ollama_timeout_seconds,
        )
    if provider != "api":
        raise WorkflowError(
            "Unsupported LLM provider. Set LINKEDIN_CAREER_MCP_LLM_PROVIDER to 'api' "
            "or 'ollama'."
        )
    if not settings.llm_api_key:
        raise WorkflowError(
            "LINKEDIN_CAREER_MCP_LLM_API_KEY is required when "
            "LINKEDIN_CAREER_MCP_LLM_PROVIDER=api."
        )
    return ApiLlmClient(
        base_url=settings.llm_api_base_url,
        model=api_model or settings.llm_api_model,
        api_key=settings.llm_api_key,
        timeout_seconds=settings.llm_api_timeout_seconds,
    )


def _build_planner_llm_client(
    settings: Settings,
    artifact_llm: ApiLlmClient | OllamaClient,
) -> ApiLlmClient | OllamaClient:
    if settings.llm_provider.casefold().strip() != "api":
        return artifact_llm
    planner_model = settings.llm_planner_api_model.strip() or settings.llm_api_model
    if planner_model == settings.llm_api_model:
        return artifact_llm
    return _build_llm_client(settings, api_model=planner_model)


async def run_from_cli(
    args: argparse.Namespace,
    *,
    settings: Settings | None = None,
    progress_callback: ProgressCallback | None = None,
) -> MatchingJobsWorkflowResult:
    settings = settings or load_settings()
    provider = LinkedInPublicJobsProvider(
        user_agent=settings.user_agent,
        timeout_seconds=settings.timeout_seconds,
    )
    service = JobSearchService(provider=provider, max_results=settings.max_results)
    llm = _build_llm_client(settings)
    planner_llm = _build_planner_llm_client(settings, llm)
    workflow = MatchingJobsWorkflow(service=service, ollama=llm, planner_llm=planner_llm)
    try:
        return await workflow.run(
            profile_dir=Path(args.profile_dir),
            blacklist_path=Path(args.blacklist_path),
            output_dir=Path(args.output_dir),
            source_resume_name=args.source_resume_name,
            current_job_description_name=args.current_job_description_name,
            location=args.location,
            date_posted=args.date_posted,
            limit_per_query=args.limit_per_query,
            max_queries=args.max_queries,
            max_jobs=args.max_jobs,
            artifact_mode=args.artifact_mode,
            cover_letter_retry_attempts=args.cover_letter_retries,
            progress_callback=progress_callback or _stderr_progress,
        )
    finally:
        await provider.aclose()
        if planner_llm is not llm:
            await planner_llm.aclose()
        await llm.aclose()


async def run_regenerate_from_cli(
    args: argparse.Namespace,
    *,
    settings: Settings | None = None,
    artifact_mode: ArtifactMode = "resumes-only",
    progress_callback: ProgressCallback | None = None,
) -> MatchingJobsWorkflowResult:
    settings = settings or load_settings()
    provider = LinkedInPublicJobsProvider(
        user_agent=settings.user_agent,
        timeout_seconds=settings.timeout_seconds,
    )
    service = JobSearchService(provider=provider, max_results=settings.max_results)
    llm = _build_llm_client(settings)
    workflow = MatchingJobsWorkflow(service=service, ollama=llm)
    try:
        return await workflow.regenerate_resumes(
            profile_dir=Path(args.profile_dir),
            output_dir=Path(args.output_dir),
            source_resume_name=args.source_resume_name,
            current_job_description_name=args.current_job_description_name,
            job_ids=args.job_ids,
            linkedin_delay_seconds=args.linkedin_delay_seconds,
            artifact_mode=artifact_mode,
            cover_letter_retry_attempts=args.cover_letter_retries,
            progress_callback=progress_callback,
        )
    finally:
        await provider.aclose()
        await llm.aclose()


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Find matching LinkedIn jobs and generate application artifacts."
    )
    parser.add_argument(
        "artifact_mode",
        nargs="?",
        default="all",
        choices=["all", "resumes-only", "cover-letters-only"],
        help="Generate resumes and cover letters by default, or only one artifact type.",
    )
    parser.add_argument("--profile-dir", default=str(DEFAULT_PROFILE_DIR))
    parser.add_argument("--blacklist-path", default=str(DEFAULT_BLACKLIST_PATH))
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--source-resume-name", default=DEFAULT_SOURCE_RESUME)
    parser.add_argument("--current-job-description-name", default=DEFAULT_CURRENT_JOB_DESCRIPTION)
    parser.add_argument("--location", default="United States")
    parser.add_argument(
        "--date-posted",
        default="past_week",
        choices=["any_time", "past_24_hours", "past_week", "past_month"],
    )
    parser.add_argument("--limit-per-query", type=int, default=10)
    parser.add_argument("--max-queries", type=int, default=6)
    parser.add_argument("--max-jobs", type=int, default=10)
    parser.add_argument(
        "--cover-letter-retries",
        type=int,
        default=1,
        help="Retry jobs still missing cover letters after the first pass. Use 0 to disable.",
    )
    return parser


def build_regenerate_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Regenerate artifacts for jobs already stored in the application database."
    )
    parser.add_argument(
        "job_ids",
        nargs="*",
        default=["all"],
        help="Use 'all', one LinkedIn job ID, space-separated IDs, or a comma-separated ID list.",
    )
    parser.add_argument("--profile-dir", default=str(DEFAULT_PROFILE_DIR))
    parser.add_argument("--output-dir", default=str(DEFAULT_OUTPUT_DIR))
    parser.add_argument("--source-resume-name", default=DEFAULT_SOURCE_RESUME)
    parser.add_argument("--current-job-description-name", default=DEFAULT_CURRENT_JOB_DESCRIPTION)
    parser.add_argument("--linkedin-delay-seconds", type=float, default=2.0)
    parser.add_argument(
        "--cover-letter-retries",
        type=int,
        default=1,
        help="Retry jobs still missing cover letters after the first pass. Use 0 to disable.",
    )
    return parser


def _llm_settings_label(settings: Settings) -> str:
    provider = settings.llm_provider.casefold().strip()
    if provider == "ollama":
        return f"ollama:{settings.ollama_model} ({settings.ollama_base_url})"
    return f"{provider}:{settings.llm_api_model} ({settings.llm_api_base_url})"


def _workflow_llm_settings_label(settings: Settings) -> str:
    provider = settings.llm_provider.casefold().strip()
    if provider != "api":
        return _llm_settings_label(settings)
    planner_model = settings.llm_planner_api_model.strip() or settings.llm_api_model
    if planner_model == settings.llm_api_model:
        return _llm_settings_label(settings)
    return (
        f"{provider}:artifact={settings.llm_api_model}; "
        f"planner={planner_model} ({settings.llm_api_base_url})"
    )


def _regenerate_processed_summary(result: MatchingJobsWorkflowResult) -> str:
    processed_count = len(result.artifacts)
    return (
        f"Jobs processed: {processed_count}/{result.jobs_found} "
        f"(resumes: {result.resumes_created}, "
        f"cover letters: {result.cover_letters_created}, "
        f"recommendations: {result.recommendations_created}, "
        f"errors: {len(result.errors)})"
    )


def _cover_letter_retry_summary(result: MatchingJobsWorkflowResult) -> str | None:
    if not result.cover_letter_retries:
        return None
    created = sum(1 for retry in result.cover_letter_retries if retry.status == "created")
    failed = sum(1 for retry in result.cover_letter_retries if retry.status == "failed")
    return (
        f"Cover letter retries: {created} created, {failed} failed "
        f"across {len(result.cover_letter_retries)} attempts."
    )


def _workplace_skip_summary(result: MatchingJobsWorkflowResult) -> list[str]:
    if not result.skipped_workplace_type:
        return []
    lines = [
        (
            f"Workplace filter skipped {len(result.skipped_workplace_type)} "
            "explicit on-site job(s)."
        )
    ]
    for label in result.skipped_workplace_type[:10]:
        lines.append(f"- {label}")
    if len(result.skipped_workplace_type) > 10:
        lines.append(f"- ...and {len(result.skipped_workplace_type) - 10} more.")
    return lines


def _experience_skip_summary(result: MatchingJobsWorkflowResult) -> list[str]:
    if not result.skipped_experience_level:
        return []
    lines = [
        (
            f"Experience filter skipped {len(result.skipped_experience_level)} "
            "internship or entry-level job(s)."
        )
    ]
    for label in result.skipped_experience_level[:10]:
        lines.append(f"- {label}")
    if len(result.skipped_experience_level) > 10:
        lines.append(f"- ...and {len(result.skipped_experience_level) - 10} more.")
    return lines


def _artifact_audit_summary(result: MatchingJobsWorkflowResult) -> list[str]:
    audit = result.artifact_audit
    lines = [
        (
            f"Artifact audit: {audit.with_resumes}/{audit.total_jobs} resumes, "
            f"{audit.with_cover_letters}/{audit.total_jobs} cover letters."
        )
    ]
    if not audit.missing_artifacts:
        lines.append("Missing artifacts: none.")
        return lines

    lines.append(
        f"Missing artifacts: {len(audit.missing_artifacts)} jobs "
        f"({audit.missing_resumes} missing resumes, "
        f"{audit.missing_cover_letters} missing cover letters)."
    )
    for artifact in audit.missing_artifacts[:10]:
        missing = ", ".join(artifact.missing)
        lines.append(
            f"- {artifact.job_id}: {artifact.company} - {artifact.title} "
            f"(missing: {missing})"
        )
    if len(audit.missing_artifacts) > 10:
        lines.append(f"- ...and {len(audit.missing_artifacts) - 10} more.")
    return lines


def _print_result_status(result: MatchingJobsWorkflowResult) -> None:
    retry_summary = _cover_letter_retry_summary(result)
    if retry_summary:
        print(retry_summary, file=sys.stderr, flush=True)
    for line in _workplace_skip_summary(result):
        print(line, file=sys.stderr, flush=True)
    for line in _experience_skip_summary(result):
        print(line, file=sys.stderr, flush=True)
    for line in _artifact_audit_summary(result):
        print(line, file=sys.stderr, flush=True)


def _build_workflow_logger(
    args: argparse.Namespace,
    *,
    operation: str,
    settings: Settings,
) -> _WorkflowRunLogger:
    logger = _WorkflowRunLogger(
        output_dir=Path(args.output_dir),
        operation=operation,
        llm_label=_workflow_llm_settings_label(settings),
    )
    print(f"Workflow log: {logger.path}", file=sys.stderr, flush=True)
    return logger


def main() -> None:
    parser = build_arg_parser()
    args = parser.parse_args()
    settings = load_settings()
    logger = _build_workflow_logger(args, operation="match-jobs", settings=settings)
    try:
        result = asyncio.run(
            run_from_cli(args, settings=settings, progress_callback=logger.progress)
        )
    except Exception as exc:
        logger.failure(exc)
        raise
    print(json.dumps(result.model_dump(mode="json"), indent=2))
    _print_result_status(result)
    logger.result(result)


def regenerate_main() -> None:
    parser = build_regenerate_arg_parser()
    args = parser.parse_args()
    settings = load_settings()
    print(f"LLM: {_llm_settings_label(settings)}", file=sys.stderr, flush=True)
    logger = _build_workflow_logger(args, operation="regenerate-resumes", settings=settings)
    try:
        result = asyncio.run(
            run_regenerate_from_cli(
                args,
                settings=settings,
                artifact_mode="resumes-only",
                progress_callback=logger.progress,
            )
        )
    except Exception as exc:
        logger.failure(exc)
        raise
    print(json.dumps(result.model_dump(mode="json"), indent=2))
    print(_regenerate_processed_summary(result), file=sys.stderr, flush=True)
    _print_result_status(result)
    logger.result(result)


def regenerate_cover_letters_main() -> None:
    parser = build_regenerate_arg_parser()
    args = parser.parse_args()
    settings = load_settings()
    print(f"LLM: {_llm_settings_label(settings)}", file=sys.stderr, flush=True)
    logger = _build_workflow_logger(
        args,
        operation="regenerate-cover-letters",
        settings=settings,
    )
    try:
        result = asyncio.run(
            run_regenerate_from_cli(
                args,
                settings=settings,
                artifact_mode="cover-letters-only",
                progress_callback=logger.progress,
            )
        )
    except Exception as exc:
        logger.failure(exc)
        raise
    print(json.dumps(result.model_dump(mode="json"), indent=2))
    print(_regenerate_processed_summary(result), file=sys.stderr, flush=True)
    _print_result_status(result)
    logger.result(result)


def regenerate_all_main() -> None:
    parser = build_regenerate_arg_parser()
    args = parser.parse_args()
    settings = load_settings()
    print(f"LLM: {_llm_settings_label(settings)}", file=sys.stderr, flush=True)
    logger = _build_workflow_logger(args, operation="regenerate-all", settings=settings)
    try:
        result = asyncio.run(
            run_regenerate_from_cli(
                args,
                settings=settings,
                artifact_mode="all",
                progress_callback=logger.progress,
            )
        )
    except Exception as exc:
        logger.failure(exc)
        raise
    print(json.dumps(result.model_dump(mode="json"), indent=2))
    print(_regenerate_processed_summary(result), file=sys.stderr, flush=True)
    _print_result_status(result)
    logger.result(result)


if __name__ == "__main__":
    main()
